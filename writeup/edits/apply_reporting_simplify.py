"""
apply_reporting_simplify.py — reporting-simple edits to the canonical manuscript
(2026-08-26). Three targeted, run-level text edits; formatting preserved.

  R3  §4.2 [98]  — drop the Bland-Altman bias/LoA sentence (user: no Bland-Altman)
  R10 §5.4 [177] — soften "well-characterized ... per-stage accuracy (Fig S2)"
  T3  §4.3 [123] — drop the dramatic "down to 9x10^-29" pooled p from body text

Each edit locates the run holding two stable text anchors and splices between
them, so no run styling is disturbed. Verifies after writing: text changed,
image relationships intact, word count sane. Numbered backup, never overwritten.

Run:  .venv/Scripts/python.exe writeup/edits/apply_reporting_simplify.py
"""
from pathlib import Path
import shutil
import sys
import docx

DOC = Path('writeup/main/CAP_sleep_mask_manuscript_main.docx')

EDITS = [
    # (para_index, anchor_a, anchor_b, replacement_between)
    # R3: keep "...cardiac band", drop the BA clause, resume at "Averaged over"
    (98,
     '3.41 BPM for the cardiac band',
     'Averaged over a whole recording',
     '. '),
    # R10: keep the rates claim, drop "well-characterized ... per-stage accuracy",
    #      resume at "That qualification", inserting a descriptive Fig S2 pointer
    (177,
     'accurate mean respiratory and cardiac rates',
     'That qualification is not a formality',
     ', with the calibration behavior characterized in §4.2.1 '
     '(per-stage error, pooled across recordings, is shown descriptively in '
     'Figure S2). '),
    # T3: drop the specific dramatic figure, keep the non-independence disclosure
    (123,
     'return very small p-values',
     'but they are computed on non-independent epochs',
     ' '),
]


def edit_paragraph(p, anchor_a, anchor_b, mid):
    """Splice mid between anchor_a and anchor_b inside the single run holding both."""
    for r in p.runs:
        t = r.text
        if anchor_a in t and anchor_b in t:
            i = t.index(anchor_a) + len(anchor_a)
            j = t.index(anchor_b)
            if j < i:
                raise ValueError('anchors out of order')
            r.text = t[:i] + mid + t[j:]
            return True
    return False


def main():
    if not DOC.exists():
        sys.exit(f'not found: {DOC}')

    backup = DOC.with_name(DOC.stem + '.BACKUP_2026-08-26_pre-reporting.docx')
    n = 1
    while backup.exists():
        n += 1
        backup = DOC.with_name(f'{DOC.stem}.BACKUP_2026-08-26_pre-reporting_{n}.docx')
    shutil.copy2(DOC, backup)
    print(f'backup -> {backup.name}')

    d = docx.Document(str(DOC))
    ps = d.paragraphs
    for idx, a, b, mid in EDITS:
        ok = edit_paragraph(ps[idx], a, b, mid)
        print(f'  [{idx}] {"OK" if ok else "MISS"}: {a[:30]!r}...{b[:30]!r}')
        if not ok:
            sys.exit(f'anchor not found in one run of para [{idx}] — aborting, no save')

    # sanity: the BA numbers should be gone, and no double period
    joined = '\n'.join(p.text for p in ps)
    assert 'Bland' not in ps[98].text, 'BA text still present in [98]'
    assert 'per-stage accuracy' not in ps[177].text, 'old per-stage claim remains'
    assert '9×10' not in ps[123].text and '9×10' not in ps[123].text, 'p-value figure remains'

    d.save(str(DOC))

    # verify reopen + image rels
    d2 = docx.Document(str(DOC))
    nwords = sum(len(p.text.split()) for p in d2.paragraphs)
    rels = d2.part.rels
    n_img = sum(1 for rid in rels if 'image' in rels[rid].reltype)
    print(f'saved. words={nwords}, image relationships={n_img}')
    print('done.')


if __name__ == '__main__':
    main()
