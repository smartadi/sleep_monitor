"""
apply_abstract_and_drop_s5.py — 2026-08-26.

Two edits to the canonical manuscript, in the reporting-simple frame:

  1. Write the abstract (was the deliberate "[Abstract - to be written.]"
     placeholder at [9]). One paragraph, reporting voice, every number quoted
     from the current Results.
  2. Drop Figure S5 (Bland-Altman) entirely — the image [246], its caption [247]
     and the explanatory paragraph [248] — completing the no-Bland-Altman pass
     (the body-text reference was already removed).

Backup is numbered, never overwritten. Verifies after write.

Run:  .venv/Scripts/python.exe writeup/edits/apply_abstract_and_drop_s5.py
"""
from pathlib import Path
import shutil
import sys
import docx

DOC = Path('writeup/main/CAP_sleep_mask_manuscript_main.docx')

ABSTRACT = (
    "This study characterizes what a dry capacitive sleep mask — three "
    "single-electrode-capacitance (SEC) sensors at the temples and forehead with "
    "an accelerometer — records across a night, against simultaneous "
    "polysomnography in six adults over twelve nights recorded at home; results "
    "are reported per recording and per subject rather than as pooled or powered "
    "claims. Breathing and the heartbeat are present in the capacitive signal "
    "throughout the night, stand 6–30 dB above the sensor noise floor, and can "
    "be followed as single continuous traces from the first hour to the last. "
    "Because the capacitive waveform produces more than one deflection per "
    "physiological event, each recording carries a scalar calibration factor k "
    "fitted against its own reference; the cardiac k of about two matches the "
    "biphasic capacitive pulse (1.70–2.43 peaks per heartbeat). After "
    "per-session calibration the mask recovers each night’s mean respiratory "
    "rate to a median of 0.24 br/min and mean cardiac rate to 1.56 BPM (per-epoch "
    "1.79 br/min and 3.41 BPM), but these are conditional on same-night "
    "calibration — without it the device does not beat a no-sensor constant — "
    "and no rate estimator follows either rate as it varies within the night. The "
    "mask does not carry cortical electrical activity: it responds at sleep "
    "spindles and delta-burst onsets, but only in low mechanical bands, and it "
    "does not track slow-wave activity (N3 discrimination at chance, 0.49, against "
    "0.74 for contact EEG through the same pipeline). Its one stage-related signal "
    "is amplitude — capacitive variance falls from wake through deep sleep and "
    "returns in REM, a depth axis consistent across subjects but too coarse to "
    "resolve individual stages. The mask is therefore a mechanical and "
    "hemodynamic monitor and a characterized overnight signal, not an EEG surrogate."
)


def main():
    if not DOC.exists():
        sys.exit(f'not found: {DOC}')
    backup = DOC.with_name(DOC.stem + '.BACKUP_2026-08-26_pre-abstract.docx')
    n = 1
    while backup.exists():
        n += 1
        backup = DOC.with_name(f'{DOC.stem}.BACKUP_2026-08-26_pre-abstract_{n}.docx')
    shutil.copy2(DOC, backup)
    print(f'backup -> {backup.name}')

    d = docx.Document(str(DOC))
    ps = d.paragraphs

    # 1. abstract
    ph = ps[9]
    assert 'Abstract' in ph.text and 'written' in ph.text, \
        f'expected placeholder at [9], found: {ph.text[:60]!r}'
    ph.runs[0].text = ABSTRACT
    for r in ph.runs[1:]:
        r.text = ''

    # 2. drop Figure S5 — caption first, verify, then image + explanation by content
    cap = ps[247]
    assert cap.text.startswith('Figure S5. Bland'), \
        f'expected S5 caption at [247], found: {cap.text[:60]!r}'
    expl = ps[248]
    assert expl.text.startswith('The limits of agreement are wide'), \
        f'expected S5 explanation at [248], found: {expl.text[:60]!r}'
    img = ps[246]  # image-only paragraph

    for p in (expl, cap, img):        # remove high index first
        p._element.getparent().remove(p._element)

    d.save(str(DOC))

    # verify
    d2 = docx.Document(str(DOC))
    full = '\n'.join(p.text for p in d2.paragraphs)
    assert 'Bland' not in full, 'Bland-Altman text still present somewhere'
    assert ABSTRACT[:40] in d2.paragraphs[9].text, 'abstract not written'
    rels = d2.part.rels
    n_img = sum(1 for rid in rels if 'image' in rels[rid].reltype)
    nwords = sum(len(p.text.split()) for p in d2.paragraphs)
    print(f'saved. abstract={len(ABSTRACT.split())} words, '
          f'total words={nwords}, image relationships={n_img}')
    print('done.')


if __name__ == '__main__':
    main()
