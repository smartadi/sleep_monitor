#!/usr/bin/env python
"""
Generate the Harmonic Ridge Prominence section as a Word document.

Methods + Results only.  Figures are embedded from writeup/figures/harmonics/.

Output: writeup/harmonics/CAP_harmonic_ridge_analysis.docx
"""

from __future__ import annotations
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).resolve().parents[2]
FIGS = ROOT / 'writeup' / 'figures' / 'harmonics'
OUT = Path(__file__).resolve().parent / 'CAP_harmonic_ridge_analysis.docx'


def setup_styles(doc: Document):
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    for level in range(1, 4):
        s = doc.styles[f'Heading {level}']
        s.font.name = 'Times New Roman'
        s.font.color.rgb = RGBColor(0, 0, 0)
        if level == 1:
            s.font.size = Pt(14)
            s.font.bold = True
            s.paragraph_format.space_before = Pt(18)
            s.paragraph_format.space_after = Pt(8)
        elif level == 2:
            s.font.size = Pt(12)
            s.font.bold = True
            s.paragraph_format.space_before = Pt(14)
            s.paragraph_format.space_after = Pt(6)
        else:
            s.font.size = Pt(11)
            s.font.bold = True
            s.font.italic = True
            s.paragraph_format.space_before = Pt(10)
            s.paragraph_format.space_after = Pt(4)


def add_figure(doc, path, caption, width=Inches(6.5)):
    if path.exists():
        doc.add_picture(str(path), width=width)
        last = doc.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        doc.add_paragraph(f'[MISSING FIGURE: {path.name}]')
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(caption)
    run.font.size = Pt(9)
    run.font.italic = True


def build():
    doc = Document()
    setup_styles(doc)

    # ── Title ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        'Spectral Ridge Prominence as a Sleep-Stage-Sensitive Feature '
        'in Capacitive Temple Sensor Signals'
    )
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.name = 'Times New Roman'

    doc.add_paragraph()

    # ════════════════════════════════════════════════════════════════
    # METHODS
    # ════════════════════════════════════════════════════════════════
    doc.add_heading('Methods', level=1)

    # ── Data ──
    doc.add_heading('Data', level=2)
    doc.add_paragraph(
        'Twelve overnight polysomnography (PSG) recordings were acquired from '
        'six healthy subjects (two nights each), totaling 81.7 hours. '
        'Capacitive temple sensors (CLE, CRE) recorded left and right '
        'temporal signals at 100 Hz. A differential channel (CH = CLE − CRE) '
        'was derived via ordinary least squares regression. '
        'Three-axis accelerometry provided motion artifact gating. '
        'AASM-scored sleep staging (Wake, N1, N2, N3, REM) from '
        'concurrent PSG served as ground truth, aligned to the capacitive '
        'recording via wall-clock timestamps.'
    )

    # ── Preprocessing ──
    doc.add_heading('Preprocessing', level=2)
    doc.add_paragraph(
        'Motion artifacts were removed from each CAP channel via OLS '
        'regression of accelerometer magnitude onto the raw signal, '
        'followed by bandpass filtering (0.05–4.0 Hz, 4th-order '
        'Butterworth, zero-phase). Per-window motion gating flagged '
        'windows where accelerometer RMS exceeded the session median '
        'plus three times the median absolute deviation (MAD).'
    )

    # ── Persistent Ridge Detection ──
    doc.add_heading('Persistent Ridge Detection', level=2)
    doc.add_paragraph(
        'Spectral ridges—narrow-band features that persist across '
        'consecutive time windows—were tracked in the power spectral '
        'density (PSD) of each CAP channel. Welch PSDs were computed '
        'in 30-second windows with 15-second overlap (10-second Welch '
        'sub-segments, 50% overlap), covering 0–5 Hz. '
        'Temporal median smoothing (9-window kernel) stabilized the '
        'PSD estimates before peak detection.'
    )
    doc.add_paragraph(
        'Spectral peaks exceeding 25% of the local median PSD were '
        'identified per window. A greedy tracker linked peaks across '
        'consecutive windows with a maximum frequency jump of 0.10 Hz '
        'and a gap tolerance of 6 windows (~90 seconds). Fragmented '
        'ridges were merged in a second pass if the gap was small and '
        'the boundary frequencies matched within twice the jump tolerance. '
        'Ridges shorter than 5 minutes or with median frequency below '
        'twice the frequency resolution were discarded.'
    )

    # ── Ridge Smoothing ──
    doc.add_heading('Ridge Frequency and Prominence Smoothing', level=2)
    doc.add_paragraph(
        'To suppress per-window jitter, each ridge’s frequency trace '
        'was smoothed with a 7-point median filter applied to the valid '
        '(non-NaN) segments. This produced flat, stable ridge lines '
        'suitable for visualization and downstream analysis.'
    )

    # ── Prominence ──
    doc.add_heading('Ridge Prominence', level=2)
    doc.add_paragraph(
        'For each ridge at each active window, prominence was defined as '
        'the ratio of the ridge’s PSD amplitude to the local spectral '
        'floor:'
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('prominence(t) = PSD_ridge(t) / median(PSD_floor(t))')
    run.font.italic = True
    run.font.size = Pt(10)

    doc.add_paragraph(
        'The local floor was computed as the median PSD in a ±0.3 Hz '
        'band around the ridge frequency, excluding a ±0.05 Hz notch '
        'centered on the peak itself. A minimum of 3 valid frequency bins '
        'in the floor region was required. Per-ridge prominence traces '
        'were smoothed with a 7-point median filter, matching the '
        'frequency trace smoothing.'
    )
    doc.add_paragraph(
        'A prominence of 1.0 indicates a ridge indistinguishable from '
        'the local spectral background; values of 3–5× indicate '
        'moderate spectral features; values exceeding 10× indicate '
        'dominant narrow-band activity. This metric is self-normalizing '
        'across time, channels, and subjects, unlike raw PSD amplitude.'
    )

    # ── Per-Window Score ──
    doc.add_heading('Per-Window Prominence Score', level=2)
    doc.add_paragraph(
        'A per-window aggregate score was computed as the maximum '
        'prominence among all active ridges at each time step, gated '
        'at a minimum prominence of 2.0× to exclude ridges barely '
        'above the spectral floor. Motion-masked windows were set to '
        'zero. The raw per-window maximum was temporally smoothed with '
        'a 15-point median filter (~3.75 minutes at 15-second step), '
        'producing a stable, slowly varying trace. The smoothed values '
        'were normalized to [0, 1] by dividing by the 95th percentile '
        'of non-zero values within each session and channel.'
    )

    # ── Visualization ──
    doc.add_heading('Visualization', level=2)
    doc.add_paragraph(
        'Per-session overlay plots were generated as six-row figures: '
        '(1) PSG hypnogram, (2–4) high-resolution spectrograms for '
        'CH, CLE, and CRE with ridge traces colored by prominence and '
        'event bars marking strong-prominence windows, (5) overlaid '
        'prominence score traces for all three channels, and '
        '(6) active ridge counts and strong-ridge counts per channel. '
        'Spectrograms were computed independently from the ridge '
        'detection PSDs using scipy.signal.spectrogram '
        '(nperseg=2048, noverlap=1920, nfft=4096 at 100 Hz), yielding '
        '~0.025 Hz frequency resolution and ~1.3 s temporal resolution '
        'with Gouraud shading for smooth interpolation. '
        'Motion-masked regions were overlaid as semi-transparent red bands.'
    )

    # ════════════════════════════════════════════════════════════════
    # RESULTS
    # ════════════════════════════════════════════════════════════════
    doc.add_heading('Results', level=1)

    # ── Ridge Counts ──
    doc.add_heading('Ridge Detection', level=2)
    doc.add_paragraph(
        'Across 12 sessions and three channels, persistent ridge '
        'detection yielded 39–95 ridges per channel per session '
        '(median ~65) after the 5-minute minimum duration filter. '
        'Ridge frequencies spanned the full 0–5 Hz analysis band, '
        'with concentrations in the respiratory (0.1–0.5 Hz) and '
        'cardiac (0.8–1.5 Hz) sub-bands. Median-filter smoothing '
        'eliminated per-window frequency jitter while preserving the '
        'overall ridge trajectory.'
    )

    # ── Prominence by Stage ──
    doc.add_heading('Prominence Discriminates Sleep Stages', level=2)
    doc.add_paragraph(
        'Pooled across all 12 sessions, ridge prominence showed '
        'clear stage dependence (Table 1). On the CLE channel, '
        'N3 epochs had the highest proportion of strong-prominence '
        'windows (27.9%), followed by N1 (22.3%) and REM (17.5%), '
        'with Wake lowest (9.5%). CRE showed a similar pattern: '
        'N3 led at 25.5%, N2 at 15.2%, and Wake at 1.8%. '
        'The differential channel CH showed more uniform distribution '
        '(N2 12.5%, N3 11.9%, N1 7.3%, Wake 7.1%).'
    )

    # Table 1
    doc.add_paragraph()
    table = doc.add_table(rows=6, cols=4, style='Table Grid')
    table.alignment = 2  # center
    headers = ['Stage', 'CH strong %', 'CLE strong %', 'CRE strong %']
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.size = Pt(9)

    data = [
        ('REM',  '0.0',  '17.5', '5.0'),
        ('N3',   '11.9', '27.9', '25.5'),
        ('N2',   '12.5', '13.5', '15.2'),
        ('N1',   '7.3',  '22.3', '7.3'),
        ('Wake', '7.1',  '9.5',  '1.8'),
    ]
    for ri, (stage, ch, cle, cre) in enumerate(data, 1):
        for ci, val in enumerate([stage, ch, cle, cre]):
            cell = table.rows[ri].cells[ci]
            cell.text = val
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(9)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        'Table 1. Percentage of windows with prominence score ≥ 0.3 '
        '(strong), by sleep stage and channel. Pooled across 12 sessions '
        '(n = 16,335 non-motion-masked windows per channel).'
    )
    run.font.size = Pt(9)
    run.font.italic = True

    # ── Per-session highlight ──
    doc.add_heading('Per-Session Examples', level=2)
    doc.add_paragraph(
        'Figure 1 shows the ridge overlay for session S1N1 '
        '(subject OS001, 7.9 hours). The CH channel exhibited '
        'strong N3 discrimination: during the 86 N3-scored epochs, '
        'median prominence reached 6.7× with 82.6% of windows '
        'classified as strong. In contrast, Wake (0.0% strong), '
        'N1 (0.0%), and REM (0.0%) showed negligible prominence on CH, '
        'while N2 contributed 25.2% strong windows concentrated in '
        'transition periods adjacent to N3 bouts. '
        'CLE and CRE showed lower but non-zero prominence across '
        'N1–N2 periods (6.7% and 9.0% strong, respectively).'
    )

    add_figure(doc, FIGS / 'ridge_overlay_S1N1.png',
               'Figure 1. Ridge overlay for S1N1 (7.9 hr). '
               'Top: hypnogram. Rows 2–4: CH, CLE, CRE spectrograms with '
               'ridge traces (colored by prominence) and strong-event bars (green). '
               'Row 5: prominence score. Row 6: active ridge counts.')

    doc.add_paragraph(
        'Session S4N2 (subject OS004, 6.0 hours) was the most '
        'ridge-rich recording, with CRE reaching 22.5% strong and '
        'CLE at 28.3%. Figure 2 shows prominent ridges concentrated '
        'in the first half of the night during sustained N2 sleep, '
        'with the prominence score trace exhibiting smooth, '
        'slowly-varying dynamics due to the temporal median filter.'
    )

    add_figure(doc, FIGS / 'ridge_overlay_S4N2.png',
               'Figure 2. Ridge overlay for S4N2 (6.0 hr). '
               'Most ridge-rich session. Strong prominence events cluster '
               'during sustained N2 in the first half of the night.')

    # ── Stage boxplot ──
    doc.add_heading('Pooled Prominence by Stage', level=2)
    doc.add_paragraph(
        'Figure 3 shows boxplots of the normalized prominence score '
        'by sleep stage across all sessions. On CLE and CRE, N3 was '
        'the only stage with substantially elevated interquartile range '
        '(IQR reaching 0.5 on CLE and 0.3 on CRE). All other stages '
        'had median scores at or near zero. CH showed flatter '
        'distributions with modest elevation during N2 and N3. '
        'The prominence metric thus captures a stage-dependent signal '
        'in the CAP data without requiring EEG or PSG labels.'
    )

    add_figure(doc, FIGS / 'ridge_overlay_score_by_stage.png',
               'Figure 3. Ridge prominence score by sleep stage, pooled '
               'across 12 sessions. N3 is the only stage with elevated '
               'prominence on CLE and CRE.')

    # ── Temporal smoothing ──
    doc.add_heading('Effect of Temporal Smoothing', level=2)
    doc.add_paragraph(
        'The combination of per-ridge median filtering (size 7, applied '
        'to both frequency and prominence traces) and per-window aggregate '
        'smoothing (size 15) produced stable, slowly varying prominence '
        'traces. Single-window spikes were suppressed while sustained '
        'prominence bouts (>3–4 minutes) were preserved. This made the '
        'score traces suitable for epoch-level sleep staging features '
        'without additional post-processing.'
    )

    # ── Prominence vs old harmonic scoring ──
    doc.add_heading('Comparison to Harmonic Ladder Scoring', level=2)
    doc.add_paragraph(
        'The prominence metric replaced an earlier harmonic ladder '
        'scoring approach that identified windows where concurrent '
        'ridges formed integer frequency ratios (f, 2f, 3f, ...). '
        'The ladder approach suffered from spurious detections: with '
        '60–90 persistent ridges per channel, coincidental integer-ratio '
        'pairs were common, inflating scores for windows with many '
        'independent ridges regardless of their physiological significance. '
        'Ridge prominence, by contrast, measures whether individual ridges '
        'genuinely stand out from the local spectral background, producing '
        'a more specific and interpretable signal.'
    )

    # ── Save ──
    doc.save(str(OUT))
    print(f'Saved -> {OUT}')


if __name__ == '__main__':
    build()
