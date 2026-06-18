#!/usr/bin/env python
"""
Build the single manuscript deliverable from DRAFT.md.

    py build_docx.py            # DRAFT.md -> CAP_sleep_mask_paper.docx
    py build_docx.py out.docx   # custom output name

DRAFT.md is the source of truth for manuscript prose. This renders it to one
clean .docx that the professor drops into the full paper. HTML comment blocks
(<!-- ... -->) in DRAFT.md are treated as internal notes and skipped.

Supported markdown: # Title / ## H1 / ### H2 / #### H3, paragraphs, '- ' bullets,
'> ' blockquote (rendered as an italic placeholder), pipe tables, and inline
**bold** / *italic*. Horizontal rules (---) and blank lines are handled.
"""
import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

HERE = Path(__file__).resolve().parent
SRC = HERE / "DRAFT.md"


def strip_comments(text: str) -> str:
    return re.sub(r"<!--.*?-->", "", text, flags=re.DOTALL)


def add_inline(paragraph, text, base_italic=False, base_color=None):
    """Render **bold** / *italic* inline spans into runs on `paragraph`."""
    # Split on bold first, then italics, keeping delimiters.
    token = re.compile(r"(\*\*.+?\*\*|\*.+?\*)")
    for part in token.split(text):
        if not part:
            continue
        bold = italic = base_italic
        if part.startswith("**") and part.endswith("**"):
            bold, part = True, part[2:-2]
        elif part.startswith("*") and part.endswith("*"):
            italic, part = True, part[1:-1]
        run = paragraph.add_run(part)
        run.bold = bold
        run.italic = italic
        if base_color is not None:
            run.font.color.rgb = base_color


def build(md: str, out: Path):
    doc = Document()
    # Sensible base style.
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    lines = md.splitlines()
    i = 0
    placeholder_color = RGBColor(0x80, 0x80, 0x80)

    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()

        if not stripped or stripped == "---":
            i += 1
            continue

        # Pipe table: collect contiguous '|' lines.
        if stripped.startswith("|"):
            block = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                block.append(lines[i].strip())
                i += 1
            rows = [
                [c.strip() for c in row.strip("|").split("|")]
                for row in block
                if not re.match(r"^\|[\s:|-]+\|?$", row)  # drop separator row
            ]
            if rows:
                table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                table.style = "Light Grid Accent 1"
                for r, cells in enumerate(rows):
                    for c, val in enumerate(cells):
                        cell = table.rows[r].cells[c]
                        cell.paragraphs[0].text = ""
                        add_inline(cell.paragraphs[0], val)
                        if r == 0:
                            for run in cell.paragraphs[0].runs:
                                run.bold = True
            continue

        # Headings.
        m = re.match(r"^(#{1,4})\s+(.*)$", stripped)
        if m:
            level = len(m.group(1))
            text = m.group(2)
            if level == 1:  # title
                p = doc.add_heading("", level=0)
                add_inline(p, text)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p = doc.add_heading("", level=level - 1)
                add_inline(p, text)
            i += 1
            continue

        # Blockquote -> italic grey placeholder.
        if stripped.startswith(">"):
            text = stripped.lstrip(">").strip()
            p = doc.add_paragraph()
            add_inline(p, text, base_italic=True, base_color=placeholder_color)
            i += 1
            continue

        # Bullet.
        if stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, stripped[2:])
            i += 1
            continue

        # Plain paragraph.
        p = doc.add_paragraph()
        add_inline(p, stripped)
        i += 1

    doc.save(str(out))
    return out


def main():
    out = Path(sys.argv[1]) if len(sys.argv) > 1 else HERE / "CAP_sleep_mask_paper.docx"
    if not out.is_absolute():
        out = HERE / out
    md = strip_comments(SRC.read_text(encoding="utf-8"))
    build(md, out)
    print(f"Built {out}")


if __name__ == "__main__":
    main()
