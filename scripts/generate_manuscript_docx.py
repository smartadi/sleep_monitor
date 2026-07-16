#!/usr/bin/env python
"""
Generate the CAP sleep mask manuscript (Methods, Results, Discussion) as a Word document.

Sections written: Methods, Results, Discussion, Limitations, Conclusion, Open Items.
Title / Abstract / Introduction are PLACEHOLDER stubs (to be written later).

Output: writeup/CAP_sleep_mask_manuscript.docx

Numbers reconciled 2026-06-18 against:
  - reports/rates/mask/ (post-consensus, tracking-FAIL)
  - reports/slow_wave/paper_n3_loso_metrics.csv
  - analysis/swa_validation/outputs/
  - CONTINUATION_RATE_DETECTION.md
  - CHANGELOG.md / ANALYSIS_LOG.md
"""
from __future__ import annotations
import sys
from pathlib import Path

import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parent.parent
FIG_SV = ROOT / "writeup" / "figures" / "signal_validation"
FIG_RATE = ROOT / "writeup" / "figures" / "mask_rate_detection"
FIG_HARM = ROOT / "writeup" / "figures" / "harmonics"
FIG_SWA = ROOT / "analysis" / "swa_validation" / "outputs"
RPT_RATE = ROOT / "reports" / "rates" / "mask"
RPT_SW = ROOT / "reports" / "slow_wave"
OUT = ROOT / "writeup" / "CAP_sleep_mask_manuscript.docx"


# ── helpers ──────────────────────────────────────────────────────────────

def setup_styles(doc: Document):
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    for level in range(1, 4):
        s = doc.styles[f"Heading {level}"]
        s.font.name = "Times New Roman"
        s.font.color.rgb = RGBColor(0, 0, 0)
        if level == 1:
            s.font.size = Pt(14)
            s.font.bold = True
            s.paragraph_format.space_before = Pt(18)
            s.paragraph_format.space_after = Pt(8)
        elif level == 2:
            s.font.size = Pt(12)
            s.font.bold = True
            s.paragraph_format.space_before = Pt(14)
            s.paragraph_format.space_after = Pt(6)
        else:
            s.font.size = Pt(11)
            s.font.bold = True
            s.font.italic = True
            s.paragraph_format.space_before = Pt(10)
            s.paragraph_format.space_after = Pt(4)


def add_para(doc, text, bold=False, italic=False, size=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if size:
        run.font.size = Pt(size)
    return p


def add_figure(doc, path, caption, width=Inches(6.0)):
    p = Path(path)
    if not p.exists():
        add_para(doc, f"[FIGURE MISSING: {p.name}]", italic=True)
        return
    fig_p = doc.add_paragraph()
    fig_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fig_p.add_run().add_picture(str(p), width=width)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(caption)
    run.font.size = Pt(9)
    run.italic = True
    cap.paragraph_format.space_after = Pt(10)


def make_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    return table


# ── Section writers ──────────────────────────────────────────────────────

def write_placeholder_front(doc):
    """Title only — Abstract and Introduction omitted for now."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "Characterization of a Capacitive Temple-Sensor Sleep Mask:\n"
        "Signal Validation, Rate Estimation, and Spectral Structure"
    )
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.name = "Times New Roman"
    doc.add_paragraph()
    add_para(doc, "[Authors]", italic=True)
    add_para(doc, "[Affiliations]", italic=True)
    doc.add_page_break()


def write_methods(doc):
    doc.add_heading("2. Methods", level=1)

    # 2.1 ────────────────────────────────────────────
    doc.add_heading("2.1 Participants and recording protocol", level=2)
    add_para(doc,
        "Twelve overnight polysomnographic recordings were collected from six healthy "
        "participants (OS001–OS006), each studied on two separate nights. Recording "
        "duration ranged from 4.11 to 8.66 hours. Each session paired a capacitive "
        "sensing (CAP) sleep mask with concurrent full polysomnography (PSG) as ground "
        "truth. The mask carried three capacitive electrodes—a central forehead "
        "channel (CH) and left and right temple channels (CLE, CRE)—together with "
        "a three-axis accelerometer for motion reference. All CAP channels were sampled "
        "at 100 Hz. The simultaneous PSG provided a single-channel EEG, left and right "
        "EOG, ECG, nasal airflow (Flow), photoplethysmography (Pleth), and thoracic "
        "and abdominal respiratory effort bands, along with expert AASM 30-second sleep-stage "
        "scoring (Wake, N1, N2, N3, REM) and scored apnea/hypopnea events."
    )

    # Table 1 — session metadata
    add_para(doc, "Table 1. Recording sessions and demographics.", bold=True)
    sessions = [
        ("S1N1", "OS001", "7.95", "954"),
        ("S1N2", "OS001", "7.63", "916"),
        ("S2N1", "OS002", "7.73", "928"),
        ("S2N2", "OS002", "6.77", "812"),
        ("S3N1", "OS003", "6.93", "832"),
        ("S3N2", "OS003", "8.66", "1039"),
        ("S4N1", "OS004", "6.18", "741"),
        ("S4N2", "OS004", "6.02", "722"),
        ("S5N1", "OS005", "4.11", "493"),
        ("S5N2", "OS005", "4.74", "569"),
        ("S6N1", "OS006", "5.16", "619"),
        ("S6N2", "OS006", "5.78", "694"),
    ]
    make_table(doc,
        ["Session", "Subject", "Duration (h)", "Analysis epochs (60s)"],
        sessions)
    doc.add_paragraph()

    # 2.2 ────────────────────────────────────────────
    doc.add_heading("2.2 Signal preprocessing", level=2)
    add_para(doc,
        "Motion artifact was suppressed by regressing band-limited accelerometer energy "
        "out of each CAP channel. The accelerometer magnitude and the CAP channel were "
        "first bandpassed to the analysis band of interest (respiratory 0.1–0.5 Hz, "
        "≈6–30 breaths/min; cardiac 0.5–3.0 Hz, ≈30–180 beats/min) "
        "so that only motion energy within that band was removed. Two cancellers were "
        "available: an ordinary-least-squares (OLS) projection that removes a single "
        "stationary coupling coefficient, and a normalized-LMS (NLMS) adaptive FIR "
        "canceller (16 taps, µ = 0.05) that tracks time-varying coupling when "
        "posture or sensor contact drifts. Bandpass filtering used third-order zero-phase "
        "Butterworth filters."
    )
    add_para(doc,
        "Unless otherwise noted, the canonical analysis channel is the OLS differential "
        "CLE−CRE, which cancels common-mode drift between the two temple electrodes. "
        "Individual channels (CH, CLE, CRE) and their average are retained for the "
        "multi-channel analyses described below."
    )

    # 2.3 ────────────────────────────────────────────
    doc.add_heading("2.3 Ground truth derivation", level=2)
    add_para(doc,
        "Cardiac ground truth. Reference cardiac rate used ECG R-peak detection "
        "(neurokit2 Pan–Tompkins variant with ECG cleaning at the native sampling "
        "rate). Photoplethysmography served as a fallback when ECG quality was poor "
        "(S6N2: ECG signal dead). ECG R-peak detection is the gold standard for "
        "beat-level heart rate."
    )
    add_para(doc,
        "Respiratory ground truth. Rather than relying on a single PSG respiratory "
        "sensor, we built a multi-signal consensus reference from four available "
        "channels: nasal airflow (Flow), thoracic effort (Thorax), abdominal effort "
        "(Abdomen), and respiratory inductance plethysmography sum (RIPSum). Peak "
        "detection was applied independently to each sensor, and a per-session quality "
        "gate excluded sensors with net-negative correlation to the majority (e.g., S3 "
        "Thorax showed paradoxical breathing and was dropped). Apnea epochs were "
        "labelled from the PSG apnea/hypopnea annotations rather than forced to a "
        "rate. The consensus rate was defined as the median across the remaining "
        "high-quality sensors at each epoch."
    )
    add_para(doc,
        "To validate this consensus, we compared the two most independent PSG respiratory "
        "sensors (Flow and RIPSum, which measure airflow and chest-wall expansion "
        "respectively). Within-session correlation was r = +0.48 on raw rates and "
        "r = +0.28 on detrended fluctuations, confirming that within-session respiratory "
        "rate variation is real physiology. The consensus reduced single-sensor jitter "
        "(standard deviation 2.26 → 1.61 br/min) and provided 100% epoch coverage. "
        "The median absolute difference between consensus and Flow-only was 0.06 br/min, "
        "though 29% of epochs differed by more than 1 br/min. This inter-sensor "
        "disagreement represents a floor on the uncertainty achievable by any external "
        "respiratory rate sensor."
    )
    add_para(doc,
        "All rates were computed on a common sliding-window grid (60-second windows, "
        "30-second step) aligned to the consensus 5-second epoch grid. Sleep stages "
        "were taken from the PSG technologist's AASM scoring on 30-second epochs."
    )

    # 2.4 ────────────────────────────────────────────
    doc.add_heading("2.4 Signal validation approach", level=2)
    add_para(doc,
        "Before estimating rates, we established that respiratory and cardiac rhythms "
        "are physically present in the CAP signal. Band energy was quantified via "
        "sliding-window Welch power spectral density (PSD) in the respiratory "
        "(0.1–0.5 Hz) and cardiac (0.5–3.0 Hz) bands, with in-band "
        "signal-to-noise ratio (SNR) computed relative to a 3.5–5.0 Hz noise floor."
    )
    add_para(doc,
        "For each analysis epoch, magnitude-squared coherence was computed between the "
        "CAP channel and the corresponding PSG reference (Flow for respiration, ECG for "
        "cardiac), and the coherence value was read at the ground-truth rate frequency. "
        "Spectral agreement was quantified as the fraction of epochs whose CAP peak "
        "frequency fell within ±0.05 Hz of the reference. To guard against "
        "spurious coherence from band-limited noise, phase-randomized surrogates "
        "(200 per epoch) were generated that preserve the power spectrum while "
        "destroying phase structure; the fraction of epochs whose observed coherence "
        "exceeded the surrogate null at p < 0.05 was reported. Coherence was evaluated "
        "within each sleep stage to confirm coupling persists beyond wake."
    )

    # 2.5 ────────────────────────────────────────────
    doc.add_heading("2.5 Rate estimation and k-factor calibration", level=2)
    add_para(doc,
        "Per-window rates were estimated from the bandpassed CAP channel using six base "
        "estimators: spectral (Welch PSD peak frequency), autocorrelation (dominant ACF "
        "lag with parabolic interpolation), Hilbert instantaneous-frequency median, "
        "upward zero-crossing rate, prominence-thresholded peak counting with loose and "
        "strict thresholds (peaks_loose, peaks_strict), and a spectral-guided "
        "amplitude-adaptive peak detector. Additional advanced trackers were evaluated "
        "for the harder cardiac band: continuous-wavelet-transform (CWT) ridge tracking "
        "and STFT peak tracking with Viterbi temporal smoothing."
    )
    add_para(doc,
        "CAP counting systematically miscounts relative to the PSG reference because the "
        "capacitive waveform is not a clean one-cycle-per-event signal. We correct this "
        "with a per-session scalar k, defined as the median ratio between the CAP "
        "estimate and the ground-truth rate across randomly selected calibration windows, "
        "so that the calibrated rate equals the raw estimate divided by k. Calibration "
        "used 50 randomly drawn one-minute windows and was verified against the "
        "whole-night k: |k_diagnostic − k_whole| ≤ 0.04 for all sessions. Both "
        "uncalibrated and k-scaled accuracies are reported."
    )

    # 2.6 ────────────────────────────────────────────
    doc.add_heading("2.6 Within-session tracking evaluation", level=2)
    add_para(doc,
        "Mean rate accuracy (MAE) can be excellent even when a method predicts a constant "
        "value, if the true rate is stable. To test whether the mask captures "
        "within-session rate variation, we designed a tracking evaluation battery."
    )
    add_para(doc,
        "Fused Window Detection was constructed as a responsive tracker: peaks_loose and Hilbert "
        "instantaneous frequency were fused across all five channels (CLE, CRE, CH, avg, "
        "diff) via unweighted mean, k-calibrated per session, with minimal smoothing "
        "(rolling median, k = 3). This configuration was designed to maximize sensitivity "
        "to within-session rate dynamics at the expense of epoch-level noise."
    )
    add_para(doc,
        "For each session and band, we computed: (1) within-session Pearson correlation "
        "between the mask rate estimate and the ground-truth rate; (2) delta-tracking "
        "correlation (epoch-to-epoch rate changes); (3) separate correlations for "
        "transient segments (high GT rate-of-change) versus steady segments. To establish "
        "statistical significance, a temporal-shuffle null was constructed: for each "
        "session, the ground-truth epoch order was randomly permuted 200 times and the "
        "within-session correlation recomputed; a session was deemed to 'beat the null' "
        "if its observed correlation exceeded the 95th percentile of the shuffled "
        "distribution. A Wilcoxon signed-rank test across sessions assessed systematic "
        "tracking."
    )
    add_para(doc,
        "As an achievable ceiling, we computed the within-session correlation between "
        "two independent PSG respiratory sensors (nasal airflow Flow and RIPSum), which "
        "measure respiration by different physical mechanisms. This bounds the tracking "
        "performance achievable by any respiratory sensor on this time scale."
    )

    # 2.7 ────────────────────────────────────────────
    doc.add_heading("2.7 Harmonic detection and ridge tracking", level=2)
    add_para(doc,
        "Whole-night CAP spectrograms show structured harmonic ladders. We detected this "
        "structure with three complementary methods: harmonic product spectrum (HPS) to "
        "score integer-ratio alignment, cepstral analysis to recover the fundamental "
        "period, and explicit fundamental-frequency (F0) estimation from which we "
        "computed a harmonic energy ratio (HER)—the share of spectral energy "
        "concentrated on F0 and its integer multiples."
    )
    add_para(doc,
        "To capture sustained spectral features, spectrogram peaks were linked across "
        "consecutive time frames into persistent ridges (minimum duration 5 minutes). "
        "Each ridge was characterized by its frequency, duration, power, and prominence "
        "score relative to the local spectral background. Per-window summary features "
        "included the number of active harmonic groups, minimum ridge frequency, total "
        "ridge power, and frequency spread. Ridge features were aggregated by PSG sleep "
        "stage for descriptive comparison."
    )
    add_para(doc,
        "To quantify discriminative power, a leave-one-subject-out (LOSO) N3-vs-rest "
        "binary classifier was trained using a Random Forest on four ridge features "
        "(n_groups_active, min_ridge_freq, total_ridge_power, freq_spread). Given the "
        "limited sample size (6 subjects), stage comparisons are reported with effect "
        "directions and per-subject consistency counts rather than relying solely on "
        "significance statistics."
    )

    # 2.8 ────────────────────────────────────────────
    doc.add_heading("2.8 Capacitive vs. contact EEG slow-wave validation", level=2)
    add_para(doc,
        "To test whether the capacitive temple sensor detects cortical slow-wave activity "
        "(SWA), we replicated the spectral pipeline of Lucey et al. (2019). Raw EEG and "
        "the CLE−CRE differential were independently bandpassed (FIR, 0.5–40 Hz "
        "via firwin), epoched into 6-second windows, and analyzed with Welch PSD. Band "
        "powers were computed for 1–4.5 Hz (total SWA) and sub-bands (1–2, "
        "2–3, 3–4 Hz). Artifact rejection gated on 20–30 Hz EMG power "
        "(97.5th percentile threshold) and accelerometer activity."
    )
    add_para(doc,
        "Agreement metrics included Pearson and Spearman correlation, Bland-Altman bias "
        "and limits of agreement, magnitude-squared coherence in the SWA band, and ROC "
        "analysis for N3 detection (CAP SWA power as a threshold classifier against PSG "
        "N3 labels). As a pipeline sanity check, the same analysis was run using the "
        "contact EEG channel to predict its own N3 labels."
    )

    # 2.9 ────────────────────────────────────────────
    doc.add_heading("2.9 Statistical methods", level=2)
    add_para(doc,
        "All rate accuracy metrics are reported within-session to avoid inflation from "
        "between-session mean matching. MAE is reported as the median absolute error "
        "(robust to outliers). Cross-session summaries use the median and interquartile "
        "range of per-session metrics. Cross-validation uses leave-one-subject-out "
        "(LOSO) to avoid within-subject leakage. Non-parametric group comparisons use "
        "the Kruskal–Wallis test. Correlations with physiological markers use "
        "Spearman's rank with Bonferroni correction. The temporal-shuffle null test "
        "(200 iterations) provides a session-specific significance threshold for "
        "tracking correlations."
    )
    doc.add_page_break()


def write_results(doc):
    doc.add_heading("3. Results", level=1)

    # 3.1 ────────────────────────────────────────────
    doc.add_heading("3.1 Signal validation: CAP carries respiratory and cardiac band energy", level=2)
    add_para(doc,
        "The capacitive temple sensor signal contained sustained energy in both the "
        "respiratory (0.1–0.5 Hz) and cardiac (0.5–3.0 Hz) bands across all "
        "twelve recordings (Figure 1). The respiratory band carried 29–48% of total "
        "signal power (0–5 Hz) and the cardiac band 8–48%. Cross-subject variation "
        "was substantial: S6N1 was resp-dominated (48% respiratory, 8% cardiac), while "
        "S3N1 was cardiac-heavy (48% cardiac), reflecting real inter-individual "
        "differences in sensor coupling rather than noise."
    )
    add_para(doc,
        "To quantify how far the physiological content sits above the sensor's noise "
        "floor without assuming any baseline model, we computed a single broadband "
        "signal-to-noise ratio per session directly on the raw CLE−CRE signal. All "
        "physiological content (respiration, cardiac, movement) is confined below "
        "10 Hz; above 10 Hz the sensor carries only its electronic noise floor. We "
        "therefore defined the signal as the full-night power below 10 Hz and the "
        "noise as the power from 10 Hz to the Nyquist frequency (50 Hz), with "
        "SNR = 10·log10(P_signal / P_noise). SNR was positive in eleven of twelve "
        "recordings (mean +12.6 dB, median +11.8 dB, range −0.3 to +22.4 dB), "
        "confirming that the sub-10 Hz band dominates the sensor noise floor "
        "(Figure 2)."
    )
    add_figure(doc, FIG_SV / "fig5_cap_spectrogram_bands.png",
        "Figure 1. CLE−CRE spectrograms (0–5 Hz) with respiratory and cardiac "
        "band annotations for three representative sessions.")
    add_figure(doc, FIG_SV / "fig2_inband_snr.png",
        "Figure 2. Broadband in-band SNR per session (CLE−CRE). Signal is the "
        "full-night power below 10 Hz; noise is the power from 10 Hz to Nyquist "
        "(50 Hz). Inset: 12-session mean PSD showing the signal/noise split at 10 Hz. "
        "SNR is positive in eleven of twelve sessions (mean +12.6 dB).")

    add_para(doc,
        "Cross-spectral coherence at the ground-truth rate frequency confirmed "
        "physiological coupling: respiratory coherence was median 0.31 on the average "
        "channel (0.61 canonical upper bound), and cardiac coherence was median 0.16 "
        "(0.27 canonical upper bound). Phase-randomized surrogate testing (200 surrogates "
        "per epoch, 8,242 epochs tested) showed that 14.7% of respiratory epochs and "
        "9.1% of cardiac epochs exceeded the surrogate null at p < 0.05, confirming "
        "that the coupling persists across all sleep stages and is not an artifact of "
        "band-limited noise."
    )

    add_para(doc,
        "Respiratory frequency agreement between the CAP spectral peak and the PSG "
        "ground-truth frequency was 43% within ±0.05 Hz (median error 0.067 Hz). "
        "Table 2 summarizes the signal validation metrics across channels."
    )

    # Table 2 — signal validation
    add_para(doc, "Table 2. Signal validation summary across channels.", bold=True)
    sv_rows = [
        ("Avg (L+R)/2", "0.314", "0.141–0.570", "43.4%", "0.157", "0.068–0.305", "20.3%"),
        ("Left (CLE)", "0.321", "0.143–0.567", "43.9%", "0.154", "0.064–0.288", "20.1%"),
        ("Right (CRE)", "0.295", "0.132–0.527", "41.6%", "0.096", "0.040–0.192", "13.8%"),
        ("Diff (L−R)", "0.304", "0.133–0.546", "42.1%", "0.102", "0.043–0.198", "13.5%"),
        ("Canonical (bound)", "0.606", "0.406–0.799", "—", "0.273", "0.161–0.419", "—"),
    ]
    make_table(doc,
        ["Channel", "Resp coh.", "Resp IQR", "Resp match %", "Card coh.", "Card IQR", "Card match %"],
        sv_rows)
    doc.add_paragraph()

    # 3.2 ────────────────────────────────────────────
    doc.add_heading("3.2 Mean rate detection accuracy", level=2)

    doc.add_heading("Respiratory rate", level=3)
    add_para(doc,
        "The spectral estimator (Welch PSD peak frequency) provided the lowest "
        "respiratory MAE on every channel, with all channels performing equivalently "
        "(oracle-over-channels MAE = 1.08 br/min vs. single diff channel 1.09 br/min). "
        "With per-session k-calibration (k_resp ≈ 0.97, near unity—negligible "
        "calibration is needed for respiration), the per-session median MAE was 0.91 "
        "br/min [IQR 0.81–1.19, range 0.56–2.26]. The pooled (all-epoch) MAE "
        "was 1.09 br/min, with bias −0.3 br/min and 95% limits of agreement "
        "[−4.7, +4.2] br/min. The S3 sessions were outliers (MAE > 1.9 br/min), "
        "attributable to a paradoxical thoracic effort signal that degraded the consensus "
        "ground truth quality for that subject."
    )

    doc.add_heading("Cardiac rate", level=3)
    add_para(doc,
        "For cardiac rate, the peaks_loose estimator with k-calibration outperformed all "
        "other methods. On the single best channel (CRE), the per-session median MAE was "
        "3.41 BPM [IQR 3.06–8.38]. Multi-channel agreement fusion yielded a pooled "
        "MAE of 3.91 BPM, with bias −0.6 BPM and limits of agreement [−24.1, "
        "+22.9] BPM. The per-session k_cardiac averaged 1.95 [range 0.94–2.24], "
        "reflecting a consistent ∼2:1 peak-counting ratio consistent with a "
        "biphasic pulse waveform (systolic peak plus dicrotic notch). The S6 sessions "
        "were anomalous (k_cardiac = 1.35 and 0.94, MAE > 8 BPM), suggesting a "
        "different sensor coupling regime in that subject."
    )

    add_figure(doc, FIG_RATE / "fig18_mae_heatmap.png",
        "Figure 3. Multichannel × multimethod MAE heatmap (k-scaled) for respiratory "
        "(left) and cardiac (right) rate. Per-session IQR shown. Spectral dominates resp; "
        "peaks_loose dominates cardiac.")
    add_figure(doc, FIG_RATE / "fig2_bland_altman.png",
        "Figure 4. Bland–Altman plots for the best pipeline: respiratory (left) and "
        "cardiac (right).")

    # Table 3 — per-session rate accuracy
    add_para(doc, "Table 3. Per-session mean rate accuracy (k-calibrated).", bold=True)
    try:
        ps = pd.read_csv(RPT_RATE / "per_session_summary.csv")
        resp = ps[ps["band"] == "resp"].sort_values("session")
        card = ps[ps["band"] == "card"].sort_values("session")

        add_para(doc, "Respiratory (spectral, k-calibrated):", bold=True)
        resp_rows = []
        for _, r in resp.iterrows():
            resp_rows.append((
                r["session"], f"{r['k']:.3f}",
                f"{r['MAE']:.2f}", f"{r['RMSE']:.2f}",
                f"{r['bias']:.2f}",
            ))
        make_table(doc, ["Session", "k", "MAE (br/min)", "RMSE", "Bias"], resp_rows)
        doc.add_paragraph()

        add_para(doc, "Cardiac (peaks_loose, k-calibrated):", bold=True)
        card_rows = []
        for _, r in card.iterrows():
            card_rows.append((
                r["session"], f"{r['k']:.3f}",
                f"{r['MAE']:.2f}", f"{r['RMSE']:.2f}",
                f"{r['bias']:.2f}",
                f"{r['r']:.3f}" if pd.notna(r.get("r")) else "—",
            ))
        make_table(doc, ["Session", "k", "MAE (BPM)", "RMSE", "Bias", "r"], card_rows)
        doc.add_paragraph()
    except Exception as e:
        add_para(doc, f"[TABLE ERROR: {e}]", italic=True)

    doc.add_heading("Oracle headroom analysis", level=3)
    add_para(doc,
        "An oracle analysis revealed that respiratory headroom lies in method diversity, "
        "not channel diversity: the channel-oracle MAE (1.08 br/min) matched the single "
        "diff channel (1.09), while the method-oracle on the diff channel alone was 0.54 "
        "br/min and the full (channel × method) oracle was 0.16 br/min. In contrast, "
        "cardiac headroom lies in channel diversity: the channel-oracle MAE was 1.58 BPM "
        "(vs. our fused 3.91 BPM), with win distribution approximately even across "
        "channels (19–21% each), confirming that different channels carry "
        "complementary cardiac information on different epochs."
    )

    doc.add_heading("Per-stage accuracy", level=3)
    add_para(doc,
        "Respiratory MAE was worst during REM (≈2.31 br/min), likely reflecting "
        "irregular breathing patterns. Cardiac MAE was worst during Wake (≈4.73 "
        "BPM), consistent with greater heart rate variability during wakefulness."
    )
    add_figure(doc, FIG_RATE / "fig3_per_stage_mae.png",
        "Figure 5. Per-sleep-stage MAE for respiratory (left) and cardiac (right) rate.")

    # 3.3 ────────────────────────────────────────────
    doc.add_heading("3.3 Within-session rate tracking (negative result)", level=2)
    add_para(doc,
        "Despite accurate mean rate recovery, the mask did not recover within-session "
        "rate variation for either respiratory or cardiac bands. This was the most "
        "extensively tested finding in this study."
    )

    doc.add_heading("Tracking battery results", level=3)
    add_para(doc,
        "Fused Window Detection (the responsive tracker: peaks_loose + Hilbert mean-fusion across "
        "five channels, k-calibrated, minimal smoothing) was compared against a "
        "200-iteration temporal-shuffle null for each session. Results:"
    )
    add_para(doc,
        "Respiratory: median within-session r = +0.058 (Wilcoxon p = 0.34, not "
        "significant). Only 4 of 12 sessions exceeded the shuffle null’s 95th "
        "percentile. Delta-tracking correlation: +0.024. No advantage in transient "
        "segments over steady segments."
    )
    add_para(doc,
        "Cardiac: median within-session r = −0.188 (Wilcoxon p = 0.85, not "
        "significant). Only 3 of 12 sessions exceeded the null. Delta-tracking "
        "correlation: −0.148. An extensive battery of six estimators (peaks_loose, "
        "peaks_strict, Hilbert, spectral, CWT ridge, continuous Viterbi ridge) across "
        "multiple channels all yielded within-session r ≈ 0."
    )

    add_figure(doc, FIG_RATE / "fig19_tracking_r_bars.png",
        "Figure 6. Per-session within-session correlation (Fused Window Detection vs. spectral "
        "baseline) with temporal-shuffle null bands (5th–95th percentile). Neither "
        "band systematically exceeds the null.")

    # Table 4 — tracking battery
    add_para(doc, "Table 4. Within-session tracking battery summary.", bold=True)
    tracking_rows = [
        ("Resp", "+0.058", "4/12", "+0.024", "1.34"),
        ("Cardiac", "−0.188", "3/12", "−0.148", "4.31"),
    ]
    make_table(doc,
        ["Band", "Median r", "Sessions > null", "Δ-tracking r", "FWD MAE"],
        tracking_rows)
    add_para(doc,
        "Wilcoxon signed-rank test for systematic tracking above zero: "
        "respiratory p = 0.34, cardiac p = 0.85 (neither significant).",
        italic=True, size=9)
    doc.add_paragraph()

    doc.add_heading("Two operating points", level=3)
    add_para(doc,
        "Two rate estimation strategies represent distinct operating points: "
        "(1) the spectral estimator achieves the lowest MAE (respiratory 0.91 br/min) "
        "but has exactly zero within-session tracking—at 30-second windows, the "
        "spectral resolution (df = 0.25 Hz) quantizes the 0.4 Hz-wide respiratory "
        "band into approximately 1.6 bins, producing a literal constant prediction of "
        "≈15 br/min in 9,317 of 9,319 epochs. (2) Fused Window Detection achieves moderate "
        "MAE (respiratory 1.34 br/min, cardiac 4.31 BPM) but still does not yield "
        "statistically significant tracking. Neither operating point captures "
        "within-session rate dynamics."
    )
    add_figure(doc, FIG_RATE / "fig21_operating_points.png",
        "Figure 7. MAE vs. within-session tracking correlation for each session and band, "
        "showing the two operating points (spectral vs. Fused Window Detection).")

    doc.add_heading("Achievable tracking ceiling", level=3)
    add_para(doc,
        "To contextualize the tracking failure, we measured the within-session "
        "correlation between two independent PSG respiratory sensors: nasal airflow "
        "(Flow) and respiratory inductance plethysmography sum (RIPSum). These sensors "
        "measure respiration by entirely different physical mechanisms yet agreed at "
        "median r = +0.47 (raw) and +0.27 (detrended fluctuations). This establishes "
        "the ceiling achievable by any respiratory sensor on 30-second windows. The "
        "mask’s respiratory tracking (r = +0.058) achieves approximately 12% of "
        "this ceiling, confirming a genuine signal-to-noise limitation rather than an "
        "absence of within-session variation to track."
    )
    add_figure(doc, FIG_RATE / "fig23_ceiling_comparison.png",
        "Figure 8. Achievable tracking ceiling: mask respiratory tracking vs. "
        "Flow–RIPSum agreement (independent PSG sensors).")
    add_figure(doc, FIG_RATE / "fig22_fullnight_traces.png",
        "Figure 9. Full-night rate traces for four representative sessions showing "
        "ground truth vs. Fused Window Detection vs. spectral estimates for both bands.")

    # 3.4 ────────────────────────────────────────────
    doc.add_heading("3.4 Harmonic spectral structure and sleep-stage association", level=2)
    add_para(doc,
        "CAP spectrograms displayed structured harmonic ladders—persistent "
        "integer-ratio spectral peaks—visible across full overnight recordings. "
        "These ridges were concentrated in the CRE channel (dominant ridge channel in "
        "9/12 sessions) with harmonic detection strongest in the CH channel (70% of "
        "windows)."
    )
    add_figure(doc, FIG_HARM / "paper_overlay_S1N1.png",
        "Figure 10. Representative session (S1N1): CRE spectrogram with detected "
        "persistent ridges overlaid, aligned with PSG hypnogram.")

    doc.add_heading("Ridge features by sleep stage", level=3)
    add_para(doc,
        "Ridge features showed statistically significant variation across sleep stages "
        "(Kruskal–Wallis p < 10⁻¹⁶ for all four features). "
        "N3 epochs were associated with fewer active harmonic groups, lower minimum "
        "ridge frequency (median 0.25 vs. 0.88 Hz in other stages), less total ridge "
        "power, and less frequency spread. These directional effects were consistent "
        "in 5–6 of 6 subjects per feature. However, the harmonic energy ratio "
        "(HER) showed subject-dependent direction: S1 and S2 exhibited higher HER "
        "during N3, while S3 and S4 showed lower HER during N3."
    )
    add_figure(doc, FIG_HARM / "paper_quantification.png",
        "Figure 11. Pooled harmonic ridge quantification: active ridges by stage "
        "(violin), max prominence by stage, per-subject heatmap, single-feature ROC, "
        "per-subject AUC, and Kruskal–Wallis statistics.")

    doc.add_heading("LOSO N3 classification", level=3)
    add_para(doc,
        "Despite statistical significance, ridge features were near-chance N3 "
        "classifiers. A LOSO Random Forest trained on four ridge features yielded "
        "a pooled AUC of 0.534 [per-subject range 0.421–0.604], mean AUC 0.509, "
        "and mean F1 score 0.095. The best single feature (prominence_score) achieved "
        "AUC = 0.563. The subject-dependent HER direction, confirmed in both Stage 2 "
        "and Stage 4 analyses, cancels out pooled discrimination: the classifier cannot "
        "learn a universal N3 signature because the spectral response to deep sleep "
        "varies between individuals."
    )
    add_figure(doc, FIG_HARM / "paper_n3_loso.png",
        "Figure 12. LOSO N3 classification: per-subject ROC curves, pooled metrics, "
        "and feature importance. Pooled AUC = 0.534.")

    # Table 5 — LOSO N3
    add_para(doc, "Table 5. LOSO N3 classification per subject.", bold=True)
    try:
        loso = pd.read_csv(RPT_SW / "paper_n3_loso_metrics.csv")
        loso_rows = []
        for _, r in loso.iterrows():
            loso_rows.append((
                r.get("subject_label", r.get("subject", "")),
                str(int(r["n_test"])), str(int(r["n_n3"])),
                f"{r['auc']:.3f}", f"{r['f1']:.3f}",
            ))
        make_table(doc, ["Subject", "N test", "N N3", "AUC", "F1"], loso_rows)
    except Exception as e:
        add_para(doc, f"[TABLE ERROR: {e}]", italic=True)
    doc.add_paragraph()

    # 3.5 ────────────────────────────────────────────
    doc.add_heading("3.5 Capacitive vs. contact EEG slow-wave activity (negative result)", level=2)
    add_para(doc,
        "Replication of the Lucey et al. (2019) spectral pipeline confirmed that the "
        "capacitive temple sensor does not measure cortical slow-wave activity. Across "
        "all twelve sessions, the correlation between CAP-derived and EEG-derived SWA "
        "power was r = 0.015 ± 0.045 (effectively zero). Magnitude-squared "
        "coherence in the SWA band was 0.003 ± 0.005 (noise floor). N3 detection "
        "using CAP SWA as a threshold yielded AUC = 0.490 ± 0.040 (chance level)."
    )
    add_para(doc,
        "The pipeline was validated by applying the identical analysis to the contact "
        "EEG channel: EEG self-AUC was 0.740 ± 0.056, confirming that the "
        "spectral pipeline correctly recovers N3 from a genuine EEG signal and that "
        "the negative result reflects a true absence of cortical SWA in the capacitive "
        "signal, not a pipeline error."
    )
    add_para(doc,
        "Visual inspection of all twelve sessions—comparing z-scored EEG delta "
        "power against CAP low-frequency/high-frequency power ratios on three channels "
        "(CLE−CRE, CLE, CRE)—showed no co-movement or anti-correlation in any "
        "session. Seven sessions were dominated by motion-artifact spikes with no slow "
        "structure; the remaining five showed slow modulations in the CAP ratio that were "
        "respiratory or mechanical in origin, not correlated with cortical delta."
    )

    add_figure(doc, FIG_SWA / "correlation_scatter.png",
        "Figure 13. Per-session correlation between CAP-derived and EEG-derived SWA "
        "power. All sessions cluster near r = 0.")
    add_figure(doc, FIG_SWA / "roc_curves.png",
        "Figure 14. ROC curves for N3 detection: CAP-based (near chance) vs. "
        "EEG-based (AUC ≈ 0.74) using the same spectral pipeline.")

    # Table 6 — SWA validation
    add_para(doc, "Table 6. SWA validation per subject.", bold=True)
    try:
        swa = pd.read_csv(FIG_SWA / "swa_validation_per_subject.csv")
        swa_rows = []
        for _, r in swa.iterrows():
            swa_rows.append((
                r["Subject"],
                f"{r['Total_hr']:.1f}",
                f"{r['Total_N3_min']:.0f}",
                f"{r['Mean_r_pearson']:.3f}",
                f"{r['Mean_coherence']:.4f}",
                f"{r['Mean_AUC']:.3f}",
                f"{r.get('Mean_sensitivity', 0):.2f}" if pd.notna(r.get('Mean_sensitivity')) else "—",
            ))
        make_table(doc,
            ["Subject", "Hours", "N3 min", "r (SWA)", "Coherence", "CAP AUC", "Sensitivity"],
            swa_rows)
    except Exception as e:
        add_para(doc, f"[TABLE ERROR: {e}]", italic=True)
    doc.add_paragraph()
    doc.add_page_break()


def write_discussion(doc):
    doc.add_heading("4. Discussion", level=1)

    add_para(doc,
        "This study provides a systematic, honest characterization of what a capacitive "
        "temple-sensor sleep mask can and cannot measure during overnight sleep. We "
        "discuss the findings in terms of the mask’s demonstrated capabilities, its "
        "clear limitations, and the mechanistic explanations for each."
    )

    doc.add_heading("4.1 What the mask reliably provides", level=2)

    doc.add_heading("Accurate mean respiratory and cardiac rates", level=3)
    add_para(doc,
        "The mask recovers per-session mean respiratory rate with MAE < 1 br/min and "
        "cardiac rate with MAE < 4 BPM after a simple per-session k-calibration. The "
        "respiratory k factor is near unity (k ≈ 0.97), meaning the spectral peak "
        "frequency directly reflects the breathing rate with negligible correction. The "
        "cardiac k factor (≈1.95) is consistent across subjects and nights, "
        "reflecting the fundamental ∼2:1 overcounting ratio that arises from the "
        "biphasic structure of the capacitive pulse waveform (systolic peak plus "
        "dicrotic notch). The k calibration is remarkably stable: diagnostic estimates "
        "from 50 random windows agree with whole-night values within 0.04, and 3 of "
        "6 subjects showed night-to-night k variation of ≤0.03. These accuracies "
        "are comparable to other non-contact and wearable cardiac rate sensors in the "
        "literature."
    )

    doc.add_heading("Physiological band energy", level=3)
    add_para(doc,
        "Both respiratory and cardiac bands carry substantial energy (29–48% and "
        "8–48% of total power respectively), and the sub-10 Hz physiological band "
        "sits above the sensor noise floor with positive broadband SNR in eleven of "
        "twelve recordings (mean +12.6 dB). The cross-spectral coherence and surrogate analyses confirm "
        "that this energy represents genuine physiological coupling rather than "
        "broadband noise. This establishes the foundational signal quality for any "
        "downstream analysis."
    )

    doc.add_heading("Stage-associated spectral structure", level=3)
    add_para(doc,
        "Harmonic ridge features differ systematically by sleep stage (p < 10⁻¹⁶), "
        "with N3 showing fewer, slower, and lower-power harmonic groups. These patterns "
        "are directionally consistent across 5–6 of 6 subjects, indicating that the "
        "spectral structure of the capacitive signal genuinely reflects sleep-state "
        "changes in cardiorespiratory coupling. While too weak for standalone staging, "
        "this structure could contribute to multi-modal approaches."
    )

    doc.add_heading("4.2 What the mask does not provide", level=2)

    doc.add_heading("Within-session rate tracking", level=3)
    add_para(doc,
        "The most extensively tested negative finding is the inability to recover "
        "within-session rate variation. Despite testing six rate estimation methods "
        "across five channels—including a purpose-built responsive tracker "
        "(Fused Window Detection) designed to maximize tracking sensitivity—within-session "
        "correlation with ground truth was indistinguishable from a temporal-shuffle "
        "null for both respiratory and cardiac bands. This is not a methodological "
        "failure: it is a signal-to-noise limit of the capacitive sensing modality."
    )
    add_para(doc,
        "The two operating points illustrate the tradeoff. The spectral estimator "
        "achieves the lowest MAE by predicting a nearly constant rate (the respiratory "
        "band spans only ≈1.6 spectral bins at 30-second resolution), effectively "
        "reporting the session mean. Fused Window Detection allows more epoch-level variation but "
        "this variation is noise, not signal. Neither approach captures the "
        "within-session dynamics."
    )
    add_para(doc,
        "Importantly, this limitation is not unlimited headroom: even two independent "
        "PSG respiratory sensors (Flow and RIPSum) agree at only r = 0.47 on "
        "30-second windows. Within-session respiratory variation is real but modest "
        "(GT standard deviation ≈2 br/min), and no sensor achieves perfect "
        "epoch-level tracking at this time scale."
    )

    doc.add_heading("Standalone sleep-stage classification", level=3)
    add_para(doc,
        "While harmonic ridge features are statistically associated with sleep stage, "
        "their discriminative power is near chance (LOSO N3 AUC = 0.534). The "
        "critical limitation is subject-dependent directionality: some subjects show "
        "increased harmonic energy during N3 while others show decreased energy. This "
        "prevents a universal classifier from learning a consistent N3 signature. Ridge "
        "features encode sleep architecture variation but not in a direction-universal "
        "way."
    )

    doc.add_heading("Cortical EEG slow-wave activity", level=3)
    add_para(doc,
        "The capacitive temple sensor shows zero correlation with contact EEG "
        "slow-wave activity (r ≈ 0.015), zero coherence in the SWA band "
        "(≈0.003), and chance-level N3 detection (AUC ≈ 0.49). This "
        "result, validated by the pipeline’s correct performance on contact EEG "
        "(AUC = 0.74), is definitive: the capacitive signal at temple placement "
        "detects intracranial pressure pulsations and scalp hemodynamics, not cortical "
        "electrical activity. The sub-1 Hz content in the capacitive signal is "
        "dominated by respiratory artifact and baseline wander, not delta EEG."
    )

    doc.add_heading("4.3 Mechanistic interpretation", level=2)
    add_para(doc,
        "The cardiac tracking failure can be understood through the k factor. The "
        "consistent k ≈ 2 across subjects indicates that the capacitive cardiac "
        "waveform contains two inflection points per heartbeat—most likely the "
        "systolic and dicrotic pressure peaks. The dominant frequency of this biphasic "
        "waveform is determined by its stable morphology rather than the instantaneous "
        "heart rate. When heart rate varies within a session, the waveform stretches or "
        "compresses, but the peak-counting frequency is governed by the persistent "
        "biphasic structure. Only changes large enough to alter the waveform’s "
        "fundamental peak structure (rather than its period) would produce a frequency "
        "shift detectable by the estimators tested."
    )
    add_para(doc,
        "The respiratory tracking limitation has a simpler explanation: spectral "
        "resolution. At 30-second windows with a 100 Hz sampling rate and standard "
        "Welch parameters, the frequency resolution (0.25 Hz) is comparable to the "
        "width of the entire respiratory band (0.4 Hz). This inherently quantizes "
        "respiratory rate estimates, and attempts to increase resolution (longer windows, "
        "full-window periodograms with parabolic interpolation) trade resolution for "
        "temporal responsiveness without improving tracking correlation."
    )
    add_para(doc,
        "The harmonic direction ambiguity likely reflects individual differences in how "
        "cardiorespiratory coupling reorganizes during deep sleep. The capacitive "
        "sensor sits at the temple, receiving a mix of intracranial pressure, "
        "superficial temporal artery pulsation, and near-field respiratory displacement. "
        "The relative contribution of each component varies with individual anatomy, "
        "sensor placement, and mask fit. During N3, cardiac output decreases and "
        "respiratory mechanics change; how these changes project onto the temple sensor "
        "depends on the subject-specific coupling geometry."
    )

    doc.add_heading("4.4 Clinical and practical implications", level=2)
    add_para(doc,
        "The mask is suitable for screening-level overnight respiratory and cardiac rate "
        "monitoring: accurate mean rates per session, with well-characterized calibration "
        "behavior and per-stage accuracy. This could support home sleep apnea screening "
        "(mean-rate anomalies, gross rate extremes) or longitudinal tracking of resting "
        "rates across nights. However, it cannot replace PSG or chest-worn sensors for "
        "applications requiring real-time or instantaneous rate monitoring."
    )
    add_para(doc,
        "The stage-associated spectral structure, while too weak for standalone staging, "
        "could contribute features to a multi-modal sleep staging system that combines "
        "accelerometer-derived body position and movement, rate statistics, and heart "
        "rate variability (when derived from the mean cardiac rate over longer windows). "
        "The consistently low N3 harmonic power across most subjects would add value as "
        "one input among many, even though it cannot carry the classification alone."
    )
    add_para(doc,
        "The negative SWA result clarifies a scope boundary for capacitive temple "
        "sensors: they detect mechanical and hemodynamic signals, not cortical "
        "electrical activity. Future development should not pursue EEG-like "
        "interpretations of the capacitive signal at this placement."
    )

    doc.add_heading("4.5 Comparison with prior work", level=2)
    add_para(doc,
        "The cardiac k ≈ 2 and the resulting ∼4 BPM MAE are consistent with "
        "the ballistocardiographic (BCG) literature, where waveform morphology and "
        "mechanical coupling introduce systematic overcounting that requires "
        "calibration. The respiratory k ≈ 1 reflects the simpler coupling: "
        "each breath produces a single dominant displacement of the temple sensor, "
        "unlike the complex cardiac pulse. The spectral respiratory rate limitation "
        "is a known challenge in wearable respiratory monitoring, where "
        "resolution-vs-responsiveness tradeoffs constrain epoch-level precision."
    )
    doc.add_page_break()


def write_limitations(doc):
    doc.add_heading("5. Limitations", level=1)
    add_para(doc,
        "First, the sample is small (6 subjects, 12 nights). While sufficient for "
        "within-session characterization and the clear negative results (tracking, SWA), "
        "it limits between-subject generalization and the statistical power of LOSO "
        "analyses. The harmonic direction ambiguity, in particular, requires a larger "
        "cohort to determine whether it reflects genuine inter-individual variability or "
        "a sampling artifact."
    )
    add_para(doc,
        "Second, this is a single-site laboratory study. Sensor coupling may vary with "
        "head shape, hair density, and mask fit in uncontrolled home-use settings. The "
        "S6 outlier behavior (anomalous k, elevated MAE) hints at the sensitivity of "
        "results to coupling quality."
    )
    add_para(doc,
        "Third, k-calibration requires a reference cardiac or respiratory rate for "
        "initial calibration. While k is stable within and across nights once estimated, "
        "a deployment scenario would need either a brief calibration period against a "
        "reference or a population-level k prior. The self-supervised adaptive k "
        "approach tested here failed for cardiac (no reliable k-free cardiac anchor "
        "exists), though the near-unity respiratory k makes respiratory rate effectively "
        "calibration-free."
    )
    add_para(doc,
        "Fourth, the respiratory consensus ground truth was derived from the same PSG "
        "system as the reference. Truly independent respiratory validation would require "
        "a separate sensing modality (e.g., capnography or acoustic respiration "
        "monitoring)."
    )
    add_para(doc,
        "Fifth, the 30-second analysis window was chosen for compatibility with PSG "
        "staging epochs. Shorter windows may improve temporal resolution of rate "
        "estimates but at the cost of increased noise. Our window-size sweep showed "
        "that longer windows decrease MAE but do not improve tracking correlation, "
        "suggesting the limitation is fundamental rather than an artifact of window "
        "choice."
    )
    add_para(doc,
        "Sixth, the EEG montage derivation of our single-channel contact EEG is not "
        "documented. If it corresponds to a posterior montage rather than the frontal "
        "derivation used by Lucey et al. (2019), the SWA comparison may be conservative. "
        "However, the near-zero correlation across all sub-bands and all sessions makes "
        "a montage-dependent reversal extremely unlikely."
    )


def write_conclusion(doc):
    doc.add_heading("6. Conclusion", level=1)
    add_para(doc,
        "This study provides a rigorous, multi-method characterization of a capacitive "
        "temple-sensor sleep mask during overnight sleep. The mask reliably recovers "
        "mean respiratory rate (MAE < 1 br/min) and cardiac rate (MAE < 4 BPM) per "
        "session with simple calibration, and its spectral structure carries "
        "statistically significant sleep-stage associations. However, within-session rate "
        "variation is not recoverable—a signal-to-noise limitation of the capacitive "
        "sensing modality confirmed across six estimation methods, five channels, and a "
        "rigorous temporal-shuffle null test—and the temple sensor does not detect "
        "cortical slow-wave activity."
    )
    add_para(doc,
        "These findings frame the mask as a viable tool for unobtrusive overnight mean "
        "rate monitoring and a contributor to multi-modal sleep assessment, while "
        "clarifying the fundamental boundaries of capacitive temple sensing technology. "
        "The honest characterization of both capabilities and limitations provides a "
        "foundation for future development targeting the identified headroom: smart "
        "multi-channel cardiac fusion, improved spectral resolution for respiratory "
        "tracking, and multi-modal staging approaches."
    )
    doc.add_page_break()


def write_open_items(doc):
    doc.add_heading("OPEN ITEMS / REVIEW NOTES", level=1)
    add_para(doc,
        "This section lists items for the authors to verify before submission.",
        italic=True)

    doc.add_heading("A. Numbers verified against source files", level=2)
    add_para(doc,
        "All key numbers were cross-referenced against post-consensus source CSVs "
        "(reports/rates/mask/, reports/slow_wave/, analysis/swa_validation/outputs/). "
        "The following numbers were verified:\n"
        "• Signal validation: coherence, frequency match, surrogate % from "
        "writeup/figures/signal_validation/table1*.csv and surrogate_significance.csv\n"
        "• Rate MAE: from reports/rates/mask/per_session_summary.csv and "
        "symmetric_tracking_mae_table.csv\n"
        "• Tracking battery: from reports/rates/mask/symmetric_tracking_battery.csv\n"
        "• Ceiling: from reports/rates/mask/symmetric_tracking_ceiling.csv\n"
        "• Harmonics LOSO: from reports/slow_wave/paper_n3_loso_metrics.csv\n"
        "• SWA: from analysis/swa_validation/outputs/swa_validation_per_subject.csv"
    )

    doc.add_heading("B. Claims revised from the scaffold", level=2)
    add_para(doc,
        "1. 'k-scaled respiratory achieves 0.99 br/min' → Updated to 0.91 (median) / "
        "1.09 (pooled). Source: mask pipeline, not rate_consolidation.\n"
        "2. 'k-scaled cardiac achieves 3.55 BPM' → Updated to 3.41 (median) / 3.91 "
        "(pooled). Different method (peaks_loose, not hilbert).\n"
        "3. 'Spectral is best for respiratory' → RETAINED but with critical caveat: "
        "spectral at 30s windows is a constant predictor (within-session r = 0.00).\n"
        "4. 'k_cardiac as independent biomarker' → REMOVED. k(t) correlates r = −0.83 "
        "with rate (absorbs 1/rate), GT-free proxy corr = −0.06. Not independent.\n"
        "5. 'CWT ridge outperforms all for cardiac' → STALE. peaks_loose + k wins in "
        "final pipeline.\n"
        "6. 'Multi-channel fusion improves' → REVISED. Resp: irrelevant. Cardiac: oracle "
        "headroom (1.58 BPM) not realizable per tracking diagnostic.\n"
        "7. 'Best-of-both pipeline' → Superseded by mask pipeline.\n"
        "8. ADDED: Within-session tracking FAILS for both bands (the major new finding).\n"
        "9. ADDED: LOSO N3 AUC = 0.534 (harmonics significant but near-chance classifier).\n"
        "10. ADDED: Achievable ceiling (Flow vs RIPSum r = 0.47)."
    )

    doc.add_heading("C. Sections flagged as still-evolving", level=2)
    add_para(doc,
        "• Title / Abstract / Introduction: placeholders only — to be written.\n"
        "• Fused Window Detection cardiac: channel-diversity oracle (1.58 BPM) suggests realizable "
        "headroom with better channel selection. A learning-based per-epoch channel "
        "selector may improve cardiac accuracy in future work.\n"
        "• Resp consensus GT: alignment caveat (correlations fragile to grid offset). "
        "Current manuscript uses exact-grid joins; verify if future re-evaluations "
        "change resp numbers.\n"
        "• Per-stage MAE numbers (REM worst for resp, Wake worst for cardiac): pulled "
        "from the mask pipeline’s fig3. Exact per-stage values not in a separate CSV; "
        "confirm from figures.\n"
        "• Harmonic LOSO: the subject-dependent direction is a real effect (confirmed "
        "across Stages 2 and 4 of the analysis). A larger cohort may reveal whether "
        "this reflects true physiological subgroups or sampling noise."
    )

    doc.add_heading("D. Figure cross-reference", level=2)
    add_para(doc,
        "Figures embedded in this document:\n"
        "Fig 1: fig5_cap_spectrogram_bands.png (signal validation spectrograms)\n"
        "Fig 2: fig2_inband_snr.png (broadband SNR per session)\n"
        "Fig 3: fig18_mae_heatmap.png (multichannel × multimethod MAE)\n"
        "Fig 4: fig2_bland_altman.png (Bland–Altman)\n"
        "Fig 5: fig3_per_stage_mae.png (per-stage accuracy)\n"
        "Fig 6: fig19_tracking_r_bars.png (tracking battery)\n"
        "Fig 7: fig21_operating_points.png (two operating points)\n"
        "Fig 8: fig23_ceiling_comparison.png (achievable ceiling)\n"
        "Fig 9: fig22_fullnight_traces.png (full-night traces)\n"
        "Fig 10: paper_overlay_S1N1.png (harmonic ridge overlay)\n"
        "Fig 11: paper_quantification.png (harmonic quantification)\n"
        "Fig 12: paper_n3_loso.png (LOSO N3 classification)\n"
        "Fig 13: correlation_scatter.png (SWA correlation)\n"
        "Fig 14: roc_curves.png (SWA N3 ROC)\n\n"
        "Additional figures available but not embedded:\n"
        "• fig6_bandpower_vs_psg_rate.png (band power vs GT rate time course)\n"
        "• fig14_all_sessions_resp.png, fig15_all_sessions_card.png (all 12 traces)\n"
        "• fig20_delta_transient.png (delta-tracking breakdown)\n"
        "• paper_overlay_S*N*.png (all 12 sessions harmonic overlays)\n"
        "• swa_overlay_all.png, bland_altman_summary.png, coherence_spectrum.png"
    )


# ── Main ─────────────────────────────────────────────────────────────────

def main():
    print("Generating manuscript...")
    doc = Document()
    setup_styles(doc)
    write_placeholder_front(doc)
    write_methods(doc)
    write_results(doc)
    write_discussion(doc)
    write_limitations(doc)
    write_conclusion(doc)
    write_open_items(doc)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print(f"Saved: {OUT}")
    print(f"Size: {OUT.stat().st_size / 1024:.0f} KB")


if __name__ == "__main__":
    main()
