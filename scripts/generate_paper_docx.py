#!/usr/bin/env python
"""
Generate the CAP sleep analysis paper as a Word document.

Currently: Signal Validation section only (Methods + Results).
Additional sections will be added incrementally.

Output: writeup/CAP_sleep_analysis_paper.docx
"""

from __future__ import annotations
import sys
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import pandas as pd

ROOT = Path(__file__).resolve().parent.parent
FIGURES = ROOT / 'writeup' / 'figures' / 'signal_validation'
RATE_FIGS = ROOT / 'writeup' / 'figures' / 'rate_accuracy'
K_FIGS = ROOT / 'writeup' / 'figures' / 'k_biomarker'
ARTIFACTS = ROOT / 'artifacts'
OUT = ROOT / 'writeup' / 'CAP_sleep_analysis_paper.docx'


def setup_styles(doc: Document):
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    for level in range(1, 4):
        sname = f'Heading {level}'
        s = doc.styles[sname]
        s.font.name = 'Times New Roman'
        s.font.color.rgb = RGBColor(0, 0, 0)
        if level == 1:
            s.font.size = Pt(14)
            s.font.bold = True
            s.paragraph_format.space_before = Pt(18)
            s.paragraph_format.space_after = Pt(8)
        elif level == 2:
            s.font.size = Pt(12)
            s.font.bold = True
            s.paragraph_format.space_before = Pt(14)
            s.paragraph_format.space_after = Pt(6)
        else:
            s.font.size = Pt(11)
            s.font.bold = True
            s.font.italic = True
            s.paragraph_format.space_before = Pt(10)
            s.paragraph_format.space_after = Pt(4)

    return doc


def add_title_page(doc: Document):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Respiratory and Cardiac Signal Extraction from '
                     'Capacitive Temple Sensors During Overnight Sleep')
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.name = 'Times New Roman'

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('[Authors]')
    run.font.size = Pt(11)
    run.font.italic = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('[Affiliations]')
    run.font.size = Pt(10)
    run.font.italic = True

    doc.add_paragraph()


def add_abstract(doc: Document):
    doc.add_heading('Abstract', level=1)
    doc.add_paragraph('[Abstract to be written after all sections are complete.]')
    doc.add_paragraph()


def add_figure(doc: Document, path: Path, caption: str, width: float = 6.0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    if path.exists():
        run.add_picture(str(path), width=Inches(width))
    else:
        p.add_run(f'[MISSING FIGURE: {path.name}]')

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.LEFT
    cap.paragraph_format.space_before = Pt(4)
    cap.paragraph_format.space_after = Pt(12)
    run = cap.add_run(caption)
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.name = 'Times New Roman'


def add_table_from_csv(doc: Document, csv_path: Path, caption: str):
    df = pd.read_csv(csv_path)

    cap = doc.add_paragraph()
    run = cap.add_run(caption)
    run.font.size = Pt(9)
    run.font.bold = True
    run.font.italic = True
    cap.paragraph_format.space_after = Pt(4)

    table = doc.add_table(rows=len(df) + 1, cols=len(df.columns))
    table.style = 'Light Shading'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for j, col in enumerate(df.columns):
        cell = table.rows[0].cells[j]
        cell.text = col
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.size = Pt(8)
                run.font.bold = True
                run.font.name = 'Times New Roman'

    for i, row in df.iterrows():
        for j, val in enumerate(row):
            cell = table.rows[i + 1].cells[j]
            cell.text = str(val)
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(8)
                    run.font.name = 'Times New Roman'

    doc.add_paragraph()


# ==============================================================================
# METHODS -- Signal Validation
# ==============================================================================

def write_methods_signal_validation(doc: Document):
    doc.add_heading('Methods and Materials', level=1)

    # 2.1 Participants
    doc.add_heading('Participants and Recording Protocol', level=2)
    doc.add_paragraph(
        'Six healthy adult volunteers (3 male, 3 female; ages 20s-60s; '
        'Pittsburgh Sleep Quality Index < 5) were recruited for two consecutive '
        'overnight polysomnography (PSG) sessions each, yielding 12 recordings '
        '(recording durations 4.11-8.66 h). Concurrent PSG included standard '
        'EEG montage, single-lead ECG, nasal airflow (Flow), finger '
        'photoplethysmography (Pleth), and thoracic respiratory effort belt. '
        'Participants simultaneously wore a prototype sleep mask instrumented with '
        'three capacitive proximity coupling (CPC) sensors positioned at the left '
        'temple (CLE), right temple (CRE), and forehead (CH), together with a '
        'three-axis accelerometer. All channels were sampled at 100 Hz and '
        'time-synchronized to the PSG recording system.'
    )

    # 2.2 Signal preprocessing
    doc.add_heading('Signal Preprocessing', level=2)
    doc.add_paragraph(
        'Motion artifacts were removed from each capacitive sensor channel by '
        'ordinary least-squares (OLS) regression of accelerometer magnitude, '
        'retaining the residual as the cleaned signal. This approach suppresses '
        'broadband motion contamination while preserving narrowband physiological '
        'components. Five channel configurations were evaluated: the three '
        'individual sensors (CLE, CRE, CH), their arithmetic average '
        '((CLE + CRE)/2, denoted Avg), and a differential signal '
        '(CLE - CRE, denoted Diff) designed to maximize common-mode '
        'rejection of non-physiological signals.'
    )
    doc.add_paragraph(
        'Bandpass filtering was applied using third-order Butterworth filters. '
        'The respiratory band was defined as 0.1-0.5 Hz (6-30 breaths/min) '
        'and the cardiac band as 0.7-4.0 Hz (42-240 BPM). '
        'The cardiac band upper limit was extended from the nominal 3.0 Hz to '
        '4.0 Hz for the signal validation analysis to fully capture higher '
        'harmonics of the ballistocardiogram (BCG) waveform. Epochs with '
        'accelerometer RMS exceeding three median absolute deviations above '
        'the session median were flagged as high-motion and excluded from '
        'spectral and coherence analyses.'
    )

    # 2.3 Ground truth
    doc.add_heading('Ground Truth Derivation', level=2)
    doc.add_paragraph(
        'Respiratory ground truth was obtained by peak detection on the nasal '
        'airflow (Flow) channel using the neurokit2 library, consistent with '
        'American Academy of Sleep Medicine (AASM) recommendations for respiratory '
        'event scoring. Cardiac ground truth was derived from ECG R-peak detection '
        'via the Pan-Tompkins algorithm (neurokit2 implementation). Both ground '
        'truth signals were computed over 30-second sliding windows with 15-second '
        'overlap, yielding per-epoch instantaneous rates. Epochs containing '
        'physiologically implausible intervals were rejected.'
    )
    doc.add_paragraph(
        'Sleep staging was performed by a certified technician following AASM '
        'criteria in 30-second epochs (Wake, N1, N2, N3, REM). PSG stage labels '
        'were used exclusively as ground truth for stratification; no PSG-derived '
        'features entered the capacitive sensor signal processing pipeline.'
    )

    # 2.4 Signal validation framework
    doc.add_heading('Signal Validation Framework', level=2)
    doc.add_paragraph(
        'Signal presence was evaluated through three complementary analyses '
        'applied to all 12 overnight recordings (8,242 motion-free epochs total).'
    )

    # Level 1
    p = doc.add_paragraph()
    run = p.add_run('Spectral peak frequency agreement. ')
    run.font.bold = True
    p.add_run(
        'For each 30-second epoch, Welch periodograms (4-second Hanning windows, '
        '50% overlap) were computed for both the bandpassed capacitive sensor signal '
        'and the corresponding PSG ground truth. The dominant spectral peak frequency '
        'was extracted within each band. Agreement was quantified as the proportion '
        'of epochs where the absolute frequency difference fell within +/-0.05 Hz '
        '(respiratory) or +/-0.15 Hz (cardiac), and as the median absolute '
        'frequency error across all epochs.'
    )

    # Level 2
    p = doc.add_paragraph()
    run = p.add_run('Magnitude-squared coherence. ')
    run.font.bold = True
    p.add_run(
        'Magnitude-squared coherence (MSC) was computed between the bandpassed '
        'capacitive sensor signal and the bandpassed PSG ground truth for each '
        '30-second epoch using Welch\'s method. Coherence was evaluated at the '
        'ground truth dominant peak frequency, providing a targeted measure of '
        'phase-locked linear coupling at the physiologically relevant frequency. '
        'This per-epoch coherence was computed for all five channel configurations '
        'and both frequency bands, then stratified by sleep stage.'
    )
    doc.add_paragraph(
        'To establish a theoretical upper bound on single-epoch coherence from '
        'the available sensor array, canonical coherence was computed. Canonical '
        'correlation analysis was applied to the multivariate capacitive sensor '
        'signal (all channels jointly) and the PSG reference, identifying the '
        'optimal linear combination of sensors that maximizes coherence with the '
        'ground truth at each epoch. This upper bound represents the best '
        'achievable performance from any linear channel fusion strategy.'
    )

    # Level 3
    p = doc.add_paragraph()
    run = p.add_run('Surrogate statistical testing. ')
    run.font.bold = True
    p.add_run(
        'To rule out spurious coupling arising from shared spectral structure '
        '(e.g., both signals having energy in the same band without true phase '
        'coupling), 200 phase-randomized surrogates were generated per epoch. '
        'Each surrogate preserved the power spectrum of the original ground truth '
        'signal but destroyed its phase relationship with the capacitive sensor '
        'signal. The cross-correlation coefficient between the real signal pair was '
        'compared against the distribution of surrogate cross-correlations. Epochs '
        'where the observed cross-correlation exceeded the 95th percentile of the '
        'surrogate distribution were classified as significant (p < 0.05).'
    )


def write_methods_rate_estimation(doc: Document):
    doc.add_heading('Rate Estimation and k-Factor Calibration', level=2)

    doc.add_paragraph(
        'Five rate estimators were applied to the bandpass-filtered CAP signal '
        'in 60-second sliding windows (5-second step): (1) spectral -- Welch PSD '
        'peak frequency; (2) autocorrelation (ACF) -- dominant lag with parabolic '
        'interpolation; (3) Hilbert -- median instantaneous frequency from the '
        'analytic signal; (4) zero-crossing -- upward zero-crossing count; '
        '(5) peaks -- local maxima counting with tunable prominence threshold. '
        'Each estimator returns a rate in Hz, converted to BPM (cardiac) or '
        'breaths/min (respiratory).'
    )
    doc.add_paragraph(
        'All five estimators systematically overcounted physiological events '
        'relative to PSG ground truth. In the cardiac band, the CAP waveform '
        'produces approximately 1.5-1.9x more detectable peaks than true '
        'heartbeats, because the ballistocardiographic impulse generates both '
        'systolic and dicrotic components per cycle. In the respiratory band, '
        'overcounting was milder (1.2-1.6x), arising from double-peaked '
        'breathing waveforms. A fixed division by 2 would overcorrect; the '
        'true ratio varies by subject and sensor coupling.'
    )
    doc.add_paragraph(
        'To correct this bias, a per-session calibration factor k was defined '
        'as the median ratio of CAP-estimated rate to PSG ground-truth rate, '
        'computed from 50 randomly sampled 1-minute windows per session: '
        'k = median(rate_CAP / rate_GT). The calibrated rate is then '
        'rate_calibrated = rate_raw / k. A diagnostic comparison confirmed that '
        'k computed from the 50-window subsample differed from the whole-session '
        'k by |delta| <= 0.04 in all 12 sessions, indicating that a brief '
        'calibration epoch suffices. The best-performing estimators were '
        'Hilbert/k for cardiac (tightest k IQR, 100% window coverage) and '
        'peaks/k for respiratory (loose detection with prominence_factor = 0.05, '
        'min_distance = 0.4 s). This k-factor approach requires a concurrent '
        'PSG reference during calibration and is therefore not self-contained; '
        'it is a bias-correction tool, not a standalone rate detector.'
    )


# ==============================================================================
# RESULTS -- Signal Validation
# ==============================================================================

def write_results_signal_validation(doc: Document):
    doc.add_heading('Results', level=1)
    doc.add_heading('Signal Presence Validation', level=2)

    # Opening paragraph
    doc.add_paragraph(
        'Before attempting rate estimation, we assessed whether respiratory and '
        'cardiac rhythms are detectable in the capacitive sensor data at all. '
        'The temple measurement site is anatomically distant from both the lungs '
        'and the heart; any physiological coupling to the capacitive sensors must '
        'propagate through bone and soft tissue, and competes with motion artifacts, '
        'sensor drift, and environmental noise. The validation below characterizes '
        'the strength and limitations of this coupling across 8,242 motion-free '
        'epochs from all 12 overnight recordings.'
    )

    # 3.1.1 Waveform and frequency evidence
    doc.add_heading('Frequency Content', level=3)
    doc.add_paragraph(
        'Visual inspection of individual epochs confirmed that the capacitive '
        'sensor signal contains oscillatory components at respiratory and cardiac '
        'frequencies. Figure 1 shows a representative 30-second N2 epoch where '
        'the bandpass-filtered CAP waveform tracks PSG nasal airflow in the '
        'respiratory band and shows rhythmic activity in the cardiac band. The '
        'power spectra of both signals share dominant peaks at matching frequencies, '
        'consistent with a common physiological source. However, such epoch-level '
        'examples are selectively chosen; the quantitative analyses below address '
        'how consistently this holds across the full dataset.'
    )

    # Figure 1
    add_figure(
        doc,
        FIGURES / 'fig1_waveform_example.png',
        'Figure 1. Representative 30-second epoch (S1N1, N2 sleep) showing '
        'bandpass-filtered capacitive sensor signal overlaid with PSG ground truth. '
        'Top: respiratory band (0.1-0.5 Hz), CAP sensor (blue) vs nasal airflow '
        '(green). Bottom: cardiac band (0.7-4.0 Hz), CAP sensor (purple) vs ECG '
        '(red). Right panels show matched power spectral density estimates.',
        width=6.2
    )

    doc.add_paragraph(
        'Across all motion-free epochs, the capacitive sensor dominant spectral '
        'peak matched the PSG respiratory frequency within +/-0.05 Hz in 43% of '
        'epochs, with a median frequency error of 0.067 Hz (Figure 2, left). '
        'The remaining 57% of epochs showed larger discrepancies, attributable to '
        'low signal-to-noise ratio, multi-peaked spectra where the wrong peak '
        'dominated, or genuine epochs where the respiratory signal was too weak to '
        'resolve. Nevertheless, the clear diagonal clustering in the scatter plot '
        'indicates that when the sensor does detect a dominant peak, it reliably '
        'tracks the true respiratory frequency.'
    )
    doc.add_paragraph(
        'Cardiac frequency agreement was considerably poorer: only 20% of epochs '
        'matched within +/-0.15 Hz, with a median error of 0.75 Hz (Figure 2, '
        'right). The scatter plot reveals a systematic pattern rather than random '
        'disagreement: the CAP sensor peak clusters near 1.0-1.5 Hz regardless '
        'of the true cardiac frequency, and many epochs show the sensor peak at '
        'approximately twice the ECG fundamental. This doubling is characteristic '
        'of the ballistocardiographic (BCG) waveform, which produces both systolic '
        'and dicrotic mechanical impulses per cardiac cycle. The BCG harmonic '
        'structure means that simple spectral peak picking will systematically '
        'overcount cardiac events -- a problem addressed by the k-factor calibration '
        'in Section 3.2.'
    )

    # Figure 2
    add_figure(
        doc,
        FIGURES / 'fig2_frequency_agreement.png',
        'Figure 2. Dominant spectral peak frequency: CAP sensor vs PSG ground '
        'truth across 8,242 motion-free epochs (Avg channel). Left: respiratory '
        'band -- 43% agreement within +/-0.05 Hz, median error 0.067 Hz. '
        'Right: cardiac band -- systematic overcounting with CAP peaks clustering '
        'near harmonics of the true cardiac frequency. Dashed line: perfect '
        'agreement.',
        width=6.2
    )

    # 3.1.2 Coherence
    doc.add_heading('Spectral Coherence', level=3)
    doc.add_paragraph(
        'Per-epoch magnitude-squared coherence, evaluated at the ground-truth '
        'peak frequency, provided a more rigorous measure of linear coupling. '
        'Respiratory coherence was modest: median 0.31 (Avg channel; IQR '
        '0.14-0.57). This indicates that the capacitive sensor and the nasal '
        'airflow signal share some phase-locked energy at the breathing frequency, '
        'but the coupling is weak -- a median coherence below 0.5 means that more '
        'than half of the variance at that frequency is unrelated between the two '
        'signals. Cardiac coherence was weaker still at 0.16 (IQR 0.07-0.31).'
    )
    doc.add_paragraph(
        'Several factors contribute to this low coherence. The capacitive sensors '
        'are positioned at the temples, anatomically distant from both the thorax '
        '(respiratory) and the heart (cardiac). The coupling path involves '
        'mechanical transmission through skull bone and soft tissue, introducing '
        'attenuation and phase distortion. For the cardiac signal specifically, '
        'the BCG waveform is inherently multi-component: systolic, diastolic, and '
        'dicrotic impulses distribute energy across multiple harmonics, so '
        'coherence evaluated at a single frequency captures only a fraction of '
        'the total cardiac-related energy. Additionally, coherence is a linear '
        'metric and will underestimate coupling that is nonlinear or intermittent.'
    )
    doc.add_paragraph(
        'Despite these low absolute values, coherence was consistently above zero '
        'across all five sleep stages (Figure 3, panels A-B). Respiratory '
        'coherence was highest during N2 and N3, when breathing is most regular '
        'and motion artifact minimal, and lowest during Wake and N1. This '
        'stage-dependent pattern is physiologically expected and argues against '
        'the coupling being purely artifactual.'
    )

    # Figure 3
    add_figure(
        doc,
        FIGURES / 'fig3_coherence_and_surrogates.png',
        'Figure 3. (A-B) Magnitude-squared coherence at the ground-truth peak '
        'frequency by sleep stage (Avg channel). Coherence is modest overall '
        '(respiratory median 0.31, cardiac 0.16) but persists across all stages. '
        '(C-D) Distribution of observed cross-correlation magnitudes; dashed red '
        'line marks the 85th percentile. Annotations show the fraction of epochs '
        'exceeding phase-randomized surrogates at p < 0.05.',
        width=6.2
    )

    # 3.1.3 Surrogate testing
    doc.add_heading('Surrogate Statistical Testing', level=3)
    doc.add_paragraph(
        'Phase-randomized surrogate testing assessed whether the observed coupling '
        'could arise from shared spectral structure alone (i.e., both signals '
        'having energy in the same band without genuine phase locking). For each '
        'of the 8,242 epochs, 200 surrogates were generated by randomizing the '
        'Fourier phases of the PSG signal while preserving its power spectrum. '
        'In the respiratory band, 14.7% of epochs exceeded the surrogate null '
        'at p < 0.05; in the cardiac band, 9.1% exceeded the null (Table 2).'
    )
    doc.add_paragraph(
        'These significance rates are modest. While they exceed the 5% expected '
        'by chance -- confirming that the coupling is statistically real at the '
        'population level -- the majority of individual epochs (85% respiratory, '
        '91% cardiac) did not reach significance. This means that in most '
        '30-second windows, the capacitive sensor signal is not distinguishable '
        'from a phase-randomized version of the PSG ground truth on a per-epoch '
        'basis. The implication is that the physiological signal is present but '
        'weak, intermittent, and embedded in substantial noise. Rate estimation '
        'from this data will therefore require approaches that aggregate '
        'information across multiple cycles within each window, rather than '
        'relying on tight phase locking to a reference.'
    )

    # 3.1.4 Channel comparison
    doc.add_heading('Channel Configuration', level=3)
    doc.add_paragraph(
        'All five channel configurations performed similarly in the respiratory '
        'band (median coherence 0.30-0.32), with the left sensor (CLE) and the '
        'average ((CLE+CRE)/2) marginally best (Table 1). For the cardiac band, '
        'CLE and Avg (0.15-0.16) outperformed CRE and the differential (0.10 '
        'each), consistent with left-dominant cardiac pulse transmission. The '
        'PCA-derived channel offered no advantage over simple averaging for either '
        'band.'
    )
    doc.add_paragraph(
        'Canonical correlation analysis provided an upper bound on the coherence '
        'achievable by any linear combination of the sensor array: 0.61 '
        '(respiratory) and 0.27 (cardiac) (Figure 4). The gap between the best '
        'single channel and this upper bound suggests that multi-channel fusion '
        'could improve signal extraction, though even the theoretical optimum '
        'remains below 0.5 for cardiac, underscoring the fundamental weakness of '
        'BCG coupling at the temple site.'
    )

    # Figure 4
    add_figure(
        doc,
        FIGURES / 'fig4_channel_comparison.png',
        'Figure 4. Median coherence at ground-truth peak frequency by channel '
        'configuration. Error bars: interquartile range. Green bar: canonical '
        'correlation upper bound (theoretical maximum from any linear sensor '
        'combination). Even the upper bound remains modest, particularly for the '
        'cardiac band.',
        width=6.2
    )

    # Table 1
    tbl_path = FIGURES / 'table1_signal_validation_summary.csv'
    if tbl_path.exists():
        add_table_from_csv(
            doc, tbl_path,
            'Table 1. Signal validation summary. Coherence at ground-truth peak '
            'frequency (median and IQR) and frequency agreement rate across '
            '8,242 motion-free epochs. All coherence values are below 0.5, '
            'indicating weak linear coupling at both measurement bands.'
        )

    # Table 2
    surr_path = FIGURES / 'surrogate_significance.csv'
    if surr_path.exists():
        add_table_from_csv(
            doc, surr_path,
            'Table 2. Surrogate significance: fraction of epochs exceeding '
            'the 95th percentile of 200 phase-randomized surrogates. Rates '
            'exceed the 5% null expectation, confirming population-level '
            'coupling, though the majority of individual epochs do not reach '
            'significance.'
        )

    # Summary paragraph
    doc.add_paragraph(
        'In summary, the capacitive temple sensors detect respiratory and cardiac '
        'rhythms, but the coupling is weak. Respiratory signals are more reliably '
        'captured (43% frequency match, 0.31 coherence) than cardiac signals '
        '(20% match, 0.16 coherence), and the cardiac band is further complicated '
        'by systematic harmonic overcounting from the BCG waveform. Per-epoch '
        'coherence and surrogate testing establish that the signals are '
        'statistically present at the population level, but are too weak and '
        'intermittent for most individual epochs to be distinguished from noise '
        'by linear metrics alone. These findings set realistic expectations for '
        'rate estimation: accurate rate extraction from this data will require '
        'methods that tolerate low per-epoch SNR and correct for the systematic '
        'overcounting bias. The rate estimation results in Section 3.2 demonstrate '
        'that this is achievable despite the weak underlying coupling.'
    )


def write_results_rate_accuracy(doc: Document):
    doc.add_heading('Rate Estimation Accuracy', level=2)

    # Cardiac
    doc.add_heading('Cardiac Rate', level=3)
    doc.add_paragraph(
        'Without k-scaling, the best cardiac estimator (ACF) produced a median '
        'MAE of 18.3 BPM across 12 sessions. The Hilbert estimator, despite '
        'higher raw error (mean 38.5 BPM due to extreme overcounting), yielded '
        'the most consistent k values (IQR 0.11-0.33) and achieved the lowest '
        'post-calibration error. After k-scaling, cardiac MAE dropped to a mean '
        'of 4.19 +/- 1.02 BPM (median 4.16, range 2.93-6.55 BPM), a 77% '
        'reduction from the ACF baseline (Table 3, Figure 5).'
    )
    doc.add_paragraph(
        'The k-factor corrects systematic bias but does not improve per-window '
        'tracking. Pearson correlation between k-scaled CAP rate and ECG ground '
        'truth remained near zero (median r = -0.09, range -0.51 to +0.10). '
        'This means the k-scaled rate is accurate on average over a session but '
        'cannot track beat-to-beat or minute-to-minute heart rate fluctuations. '
        'The residual MAE of 4.2 BPM reflects irreducible noise: the weak '
        'BCG coupling at the temple (coherence 0.16, Section 3.1) limits how '
        'precisely any single-channel method can estimate instantaneous rate.'
    )
    doc.add_paragraph(
        'Cardiac k ranged from 1.48 (S6N2) to 1.93 (S2N1) across sessions '
        '(Table 3). Within subjects, night-to-night k was stable: 3 of 6 '
        'subjects had |delta_k| <= 0.03, with the largest shift at 0.19 (OS002). '
        'The 50-window diagnostic k matched the whole-session k to within 0.04 '
        'in every case, confirming that a short calibration period is sufficient.'
    )

    # Cardiac table
    card_csv = ARTIFACTS / 'hilbert_scaled_per_session.csv'
    if card_csv.exists():
        df = pd.read_csv(card_csv)
        tbl_df = df[['session', 'subject', 'duration_hr', 'k_whole',
                      'mae_acf', 'mae_kwhole', 'rmse_kwhole',
                      'bias_kwhole', 'r_kwhole']].copy()
        tbl_df.columns = ['Session', 'Subject', 'Duration (h)', 'k',
                          'MAE baseline (BPM)', 'MAE scaled (BPM)',
                          'RMSE (BPM)', 'Bias (BPM)', 'r']
        for c in ['Duration (h)', 'k', 'MAE baseline (BPM)', 'MAE scaled (BPM)',
                   'RMSE (BPM)', 'Bias (BPM)', 'r']:
            tbl_df[c] = tbl_df[c].apply(lambda x: f'{x:.2f}')
        tbl_df.to_csv(ROOT / 'writeup' / 'figures' / 'rate_accuracy' /
                       'table3_cardiac_per_session.csv', index=False)
        add_table_from_csv(
            doc,
            ROOT / 'writeup' / 'figures' / 'rate_accuracy' /
            'table3_cardiac_per_session.csv',
            'Table 3. Cardiac rate accuracy per session. k: calibration factor '
            '(Hilbert/k method). MAE baseline: uncalibrated ACF estimator. '
            'MAE scaled: Hilbert rate divided by k. r: Pearson correlation '
            'between scaled rate and ECG ground truth.'
        )

    # Cardiac MAE figure
    add_figure(
        doc,
        RATE_FIGS / 'all_sessions_cardiac_MAE.png',
        'Figure 5. Per-session cardiac MAE for three estimators: ACF baseline '
        '(gray), raw Hilbert (orange), and k-scaled Hilbert (red). k values '
        'annotated per session. k-scaling reduces MAE from 10-25 BPM to 3-7 BPM '
        'across all 12 sessions.',
        width=6.2
    )

    # Respiratory
    doc.add_heading('Respiratory Rate', level=3)
    doc.add_paragraph(
        'Respiratory rate estimation used peak counting with loose detection '
        'parameters (prominence_factor = 0.05, min_distance = 0.4 s) followed '
        'by k-scaling. Baseline MAE (unscaled) was 3.09 breaths/min (mean '
        'across 12 sessions). After k-scaling, MAE decreased to 2.20 breaths/min, '
        'a 25% reduction. Improvement was not uniform: 11 of 12 sessions improved, '
        'but S6N2 showed a marginal increase (Table 4, Figure 6).'
    )
    doc.add_paragraph(
        'Respiratory k ranged from 1.12 (S2N2) to 1.62 (S3N2). Unlike cardiac k, '
        'respiratory k showed a clear subject-level clustering: OS003 and OS004 '
        'had k near 1.4-1.6 (consistent double-peaked breathing waveform), while '
        'OS001 and OS002 had k near 1.2 (simpler waveform). Per-window correlation '
        'between scaled rate and Flow ground truth was weak (median r = 0.16, '
        'range -0.02 to 0.43), again indicating bias correction without '
        'per-window precision. The loose peak detection deliberately trades '
        'per-window accuracy for lower systematic bias; tighter detection '
        'parameters improved correlation but increased MAE.'
    )

    # Respiratory table
    resp_csv = ARTIFACTS / 'peak_ratio_per_session.csv'
    if resp_csv.exists():
        df = pd.read_csv(resp_csv)
        tbl_df = df[['session', 'subject', 'duration_hr', 'k_whole',
                      'mae_base', 'mae_kwhole', 'rmse_kwhole',
                      'bias_kwhole', 'r_kwhole']].copy()
        tbl_df.columns = ['Session', 'Subject', 'Duration (h)', 'k',
                          'MAE baseline', 'MAE scaled',
                          'RMSE', 'Bias', 'r']
        for c in ['Duration (h)', 'k', 'MAE baseline', 'MAE scaled',
                   'RMSE', 'Bias', 'r']:
            tbl_df[c] = tbl_df[c].apply(lambda x: f'{x:.2f}')
        tbl_df.to_csv(ROOT / 'writeup' / 'figures' / 'rate_accuracy' /
                       'table4_resp_per_session.csv', index=False)
        add_table_from_csv(
            doc,
            ROOT / 'writeup' / 'figures' / 'rate_accuracy' /
            'table4_resp_per_session.csv',
            'Table 4. Respiratory rate accuracy per session. k: calibration '
            'factor (peaks/k method). MAE in breaths/min. r: Pearson '
            'correlation between scaled rate and nasal airflow ground truth.'
        )

    # Respiratory MAE figure
    add_figure(
        doc,
        RATE_FIGS / 'all_sessions_resp_MAE.png',
        'Figure 6. Per-session respiratory MAE: baseline raw peaks (gray) vs '
        'k-scaled peaks (green). k values annotated. Improvement is consistent '
        'but modest (25% mean reduction).',
        width=6.2
    )

    # Summary
    doc.add_paragraph(
        'The k-factor approach reduces systematic overcounting bias in both bands '
        'but has two fundamental limitations. First, it requires a PSG reference '
        'for calibration and is therefore not a standalone rate detector. Second, '
        'it corrects the mean rate over a session without improving window-level '
        'tracking (r near zero in both bands). The residual error reflects the '
        'inherent noise floor of single-channel capacitive sensing at the temple: '
        'the weak, intermittent coupling documented in Section 3.1 limits '
        'instantaneous rate precision regardless of the estimation method used.'
    )


def write_results_k_biomarker(doc: Document):
    doc.add_heading('k-Factor as a Physiological Variable', level=2)

    doc.add_paragraph(
        'Although k was introduced as a calibration constant, computing it per '
        'sliding window across each session revealed that k_cardiac varies '
        'systematically with sleep stage. Kruskal-Wallis testing across five '
        'stages yielded H = 609, p = 1.6e-130 (all sessions pooled). Stage '
        'medians were: N1 = 1.71, N2 = 1.65, N3 = 1.65, Wake = 1.61, '
        'REM = 1.58 (Figure 7). The ordering -- N1 highest, REM lowest -- '
        'is consistent with BCG waveform complexity: N1 sleep is characterized '
        'by autonomic instability and frequent arousals, producing irregular '
        'BCG morphology with more detectable sub-peaks per heartbeat, while '
        'REM is accompanied by muscle atonia and more stereotyped hemodynamics.'
    )
    doc.add_paragraph(
        'k_cardiac correlated weakly with established PSG biomarkers (Figure 8): '
        'SDNN (r = -0.25), EEG delta power (r = -0.16), and accelerometer RMS '
        '(r = +0.16), all p < 1e-4 after Bonferroni correction. The negative '
        'SDNN correlation indicates that higher HRV is associated with simpler '
        'BCG waveforms (fewer overcounted peaks), which is physiologically '
        'plausible but the effect size is small. The autocorrelation halflife '
        'of k_cardiac was 1.4 minutes, indicating that waveform complexity '
        'changes slowly, on the timescale of sleep stage transitions rather '
        'than individual heartbeats.'
    )
    doc.add_paragraph(
        'k_resp showed no comparable stage dependence. Its strongest correlate '
        'was accelerometer RMS (r = +0.29), suggesting that respiratory k '
        'reflects motion-related noise contamination rather than physiology. '
        'One session (S6N2) had an anomalous k_cardiac of 0.79, indicating '
        'the Hilbert estimator was undercounting rather than overcounting, '
        'likely due to poor sensor contact. These results suggest that '
        'k_cardiac captures a real, if weak, physiological signal related to '
        'BCG waveform complexity, but the effect size is too small for k alone '
        'to serve as a sleep stage discriminator.'
    )

    # k by stage figure
    add_figure(
        doc,
        K_FIGS / 'pooled_k_by_sleep_stage.png',
        'Figure 7. k by sleep stage (all sessions pooled). Left: k_resp shows '
        'no clear stage ordering. Right: k_cardiac shows N1 > N2/N3 > Wake > '
        'REM, consistent with BCG complexity varying with autonomic tone. '
        'Kruskal-Wallis p < 1e-130 (cardiac).',
        width=6.2
    )

    # k correlation heatmap
    add_figure(
        doc,
        K_FIGS / 'pooled_correlation_heatmap.png',
        'Figure 8. Spearman correlations between per-window k and PSG '
        'biomarkers (all sessions pooled). k_cardiac shows weak negative '
        'correlations with SDNN and delta power. k_resp correlates primarily '
        'with accelerometer RMS (motion artifact). All |r| < 0.3.',
        width=5.0
    )


# ==============================================================================
# Main
# ==============================================================================

def main():
    doc = Document()
    setup_styles(doc)

    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    add_title_page(doc)
    add_abstract(doc)

    # Methods
    write_methods_signal_validation(doc)
    write_methods_rate_estimation(doc)

    # Results
    write_results_signal_validation(doc)
    write_results_rate_accuracy(doc)
    write_results_k_biomarker(doc)

    # Placeholder for future sections
    doc.add_page_break()
    doc.add_heading('Discussion', level=1)
    doc.add_paragraph(
        '[Discussion sections to be added: k-factor approach, '
        'k_cardiac as biomarker, comparison with literature.]'
    )

    doc.save(str(OUT))
    print(f'Paper saved to {OUT}')
    print(f'  Methods: 2.1-2.5')
    print(f'  Results: 3.1 (Signal Validation), 3.2 (Rate Accuracy), '
          f'3.3 (k Biomarker)')
    print(f'  Figures: 8, Tables: 4')


if __name__ == '__main__':
    main()
