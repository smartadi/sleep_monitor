#!/usr/bin/env python
"""
Generate the rate estimation section as a standalone Word document.

Covers:
  1. Original multi-channel consolidation (no-k methods, CWT, Viterbi)
  2. Hybrid pipeline (adaptive peaks, Kalman tracker, best-of-both)

Output: writeup/CAP_rate_consolidation_section.docx
"""

from __future__ import annotations
import sys
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import pandas as pd
import numpy as np

ROOT = Path(__file__).resolve().parent.parent
FIGURES = ROOT / 'writeup' / 'figures' / 'rate_consolidation'
BEST_DIR = ROOT / 'reports' / 'rates' / 'best_pipeline'
PHASE3_DIR = ROOT / 'reports' / 'rates' / 'hybrid_phase3'
PHASE2_DIR = ROOT / 'reports' / 'rates' / 'hybrid_phase2'
PHASE1_DIR = ROOT / 'reports' / 'rates' / 'hybrid_phase1'
OUT = ROOT / 'writeup' / 'CAP_rate_consolidation_section.docx'

SESSIONS = [
    'S1N1', 'S1N2', 'S2N1', 'S2N2', 'S3N1', 'S3N2',
    'S4N1', 'S4N2', 'S5N1', 'S5N2', 'S6N1', 'S6N2',
]


def setup_styles(doc: Document):
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    for level in range(1, 4):
        sname = f'Heading {level}'
        s = doc.styles[sname]
        s.font.name = 'Times New Roman'
        s.font.color.rgb = RGBColor(0, 0, 0)
        if level == 1:
            s.font.size = Pt(14)
            s.font.bold = True
            s.paragraph_format.space_before = Pt(18)
            s.paragraph_format.space_after = Pt(8)
        elif level == 2:
            s.font.size = Pt(12)
            s.font.bold = True
            s.paragraph_format.space_before = Pt(14)
            s.paragraph_format.space_after = Pt(6)
        else:
            s.font.size = Pt(11)
            s.font.bold = True
            s.font.italic = True
            s.paragraph_format.space_before = Pt(10)
            s.paragraph_format.space_after = Pt(4)

    return doc


def add_figure(doc: Document, path: Path, caption: str, width: float = 6.0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    if path.exists():
        run.add_picture(str(path), width=Inches(width))
    else:
        p.add_run(f'[MISSING FIGURE: {path.name}]')

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.LEFT
    cap.paragraph_format.space_before = Pt(4)
    cap.paragraph_format.space_after = Pt(12)
    run = cap.add_run(caption)
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.name = 'Times New Roman'


def add_table(doc: Document, df: pd.DataFrame, caption: str):
    cap = doc.add_paragraph()
    run = cap.add_run(caption)
    run.font.size = Pt(9)
    run.font.bold = True
    run.font.italic = True
    cap.paragraph_format.space_after = Pt(4)

    table = doc.add_table(rows=len(df) + 1, cols=len(df.columns))
    table.style = 'Light Shading'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for j, col in enumerate(df.columns):
        cell = table.rows[0].cells[j]
        cell.text = col
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.size = Pt(8)
                run.font.bold = True
                run.font.name = 'Times New Roman'

    for row_idx, (i, row) in enumerate(df.iterrows()):
        for j, val in enumerate(row):
            cell = table.rows[row_idx + 1].cells[j]
            if isinstance(val, float):
                if np.isnan(val):
                    cell.text = '--'
                elif abs(val) >= 0.1:
                    cell.text = f'{val:.2f}'
                else:
                    cell.text = f'{val:.3f}'
            else:
                cell.text = str(val)
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(8)
                    run.font.name = 'Times New Roman'

    doc.add_paragraph()


# ==============================================================================
# PART 1: ORIGINAL CONSOLIDATION
# ==============================================================================

def write_consolidation_methods(doc: Document):
    doc.add_heading('Part 1: Multi-Channel Rate Estimation Without PSG Calibration', level=1)
    doc.add_heading('Methods', level=2)

    doc.add_paragraph(
        'The k-factor rate estimation pipeline described previously requires a '
        'per-session calibration step using PSG ground truth. This is adequate '
        'for laboratory validation but presents a problem for deployment: if the '
        'device needs PSG to calibrate itself, it cannot function independently. '
        'We therefore evaluated whether acceptable rate accuracy could be achieved '
        'without any PSG-derived correction, using only the raw sensor signals and '
        'a multi-channel fusion strategy.'
    )

    doc.add_heading('Method and Channel Benchmark', level=3)
    doc.add_paragraph(
        'Five rate estimation methods were evaluated on five channel '
        'configurations across all 12 overnight recordings. The methods were: '
        '(1) Welch periodogram peak frequency, (2) autocorrelation function '
        '(ACF) lag at the first prominent peak, (3) Hilbert instantaneous '
        'frequency, (4) zero-crossing rate, and (5) local maxima counting. The '
        'channel configurations were the three individual sensors (CLE, CRE, CH), '
        'their arithmetic average (Avg = (CLE + CRE) / 2), and a left-right '
        'differential (Diff = CLE - CRE). Each method was applied to '
        'non-overlapping 30-second windows in both the respiratory (0.1-0.5 Hz) '
        'and cardiac (0.7-4.0 Hz) bands. No k-factor scaling was applied.'
    )

    doc.add_heading('Channel Fusion', level=3)
    doc.add_paragraph(
        'Two fusion strategies were tested. In confidence-weighted fusion, each '
        'channel contributes its rate estimate weighted by a composite signal '
        'quality score (SNR, spectral concentration, ACF prominence). In '
        'agreement-filtered fusion, channels are first screened for cross-method '
        'consistency before confidence-weighted averaging. An oracle selector '
        'picks the best channel per window with hindsight.'
    )

    doc.add_heading('CWT Ridge Tracking', level=3)
    doc.add_paragraph(
        'For cardiac specifically, a CWT ridge method was added: Morlet wavelet '
        'scalogram over 32 log-spaced scales, extracting the dominant ridge '
        'frequency per window. The motivation was that CWT may resolve the '
        'fundamental cardiac rate directly from the time-frequency representation.'
    )

    doc.add_heading('Temporal Smoothing', level=3)
    doc.add_paragraph(
        'All per-window estimates were post-processed with a Viterbi-style '
        'temporal smoother that penalizes epoch-to-epoch jumps exceeding '
        'physiological limits (5 BPM/epoch cardiac, 2 br/min/epoch respiratory).'
    )


def write_consolidation_results(doc: Document):
    doc.add_heading('Results', level=2)

    doc.add_heading('Method Benchmark', level=3)
    doc.add_paragraph(
        'For respiratory rate, spectral (Welch) was the clear winner at 1.5 br/min '
        'MAE, outperforming even k-scaled peak counting (2.20 br/min). For cardiac, '
        'no method performed well without k: ACF was best at 12.5 BPM, with Hilbert '
        '(38 BPM) and peak counting (46-51 BPM) severely misled by BCG harmonics.'
    )

    add_figure(doc, FIGURES / 'phase1_method_channel_heatmap.png',
        'Figure 1. MAE for five methods across five channels, no k-scaling. '
        'Left: respiratory (best: spectral 1.5 br/min). Right: cardiac '
        '(best: ACF 12.5 BPM).', width=6.2)

    doc.add_heading('Channel Fusion', level=3)
    doc.add_paragraph(
        'For respiratory, fusion provided no benefit (all strategies 1.5 br/min). '
        'For cardiac, confidence-weighted fusion reduced MAE by 19% to 10.1 BPM. '
        'The oracle (4.5 BPM) shows much greater potential with better channel selection.'
    )

    add_figure(doc, FIGURES / 'phase2_fusion_comparison.png',
        'Figure 2. Multi-channel fusion comparison. Respiratory: all strategies '
        'equivalent. Cardiac: confidence-weighted 10.1 BPM, oracle 4.5 BPM.',
        width=6.2)

    doc.add_heading('CWT and Viterbi Smoothing', level=3)
    doc.add_paragraph(
        'CWT achieved 11.6 BPM cardiac MAE on Avg channel -- best without k. '
        'Viterbi smoothing reduced cardiac MAE by 31-35%, yielding 6.6 BPM for '
        'the best pipeline (confidence-weighted + Viterbi). Respiratory was '
        'unchanged by smoothing.'
    )

    add_figure(doc, FIGURES / 'phase4_viterbi_improvement.png',
        'Figure 3. Viterbi temporal smoothing effect. Respiratory: unchanged. '
        'Cardiac: 31-35% reduction. Best: fused+Viterbi 6.6 BPM.',
        width=6.2)

    doc.add_heading('Agreement and Per-Stage Accuracy', level=3)

    add_figure(doc, FIGURES / 'phase5_bland_altman.png',
        'Figure 4. Bland-Altman: fused+Viterbi vs PSG. Respiratory bias +0.2 '
        'br/min, LoA [-5.0, +5.5]. Cardiac bias -7.9 BPM, LoA [-44.8, +28.9].',
        width=6.0)

    add_figure(doc, FIGURES / 'phase5_per_stage_mae.png',
        'Figure 5. Per-stage MAE. Respiratory stable at 1.5-1.6 br/min '
        '(REM: 2.5). Cardiac best in N1 (5.5 BPM), worst in REM (7.9 BPM).',
        width=6.2)


# ==============================================================================
# PART 2: HYBRID PIPELINE
# ==============================================================================

def write_hybrid_methods(doc: Document):
    doc.add_heading('Part 2: Hybrid Rate Pipeline', level=1)
    doc.add_heading('Methods', level=2)

    doc.add_paragraph(
        'The consolidation analysis revealed that the optimal rate estimation '
        'strategy differs between bands. For respiratory rate, the Welch spectral '
        'method already achieves excellent accuracy without PSG calibration. For '
        'cardiac rate, the Hilbert instantaneous frequency method with per-session '
        'k-scaling remains superior to all PSG-free approaches. We therefore '
        'developed a hybrid pipeline that combines the best approach for each '
        'band, adds adaptive peak detection and Kalman filtering for respiratory, '
        'and applies multi-channel fusion and temporal smoothing to both.'
    )

    doc.add_heading('Adaptive Peak Detector', level=3)
    doc.add_paragraph(
        'A new adaptive peak detector (rate_adaptive_peaks) was developed to '
        'improve on fixed-parameter peak counting. Three innovations were '
        'introduced: (1) spectral guidance -- the dominant Welch PSD frequency '
        'sets the expected rate, and minimum peak distance is set to 0.6x the '
        'expected period rather than using the band ceiling; (2) amplitude-adaptive '
        'prominence -- a rolling median absolute deviation (MAD) of the signal '
        'envelope replaces the fixed standard deviation threshold, adapting to '
        'local amplitude drift over the night; (3) inter-peak-interval (IPI) '
        'validation -- the coefficient of variation of detected intervals is '
        'checked against a threshold (CV < 0.40), with MAD-based outlier rejection '
        'of anomalous intervals before rate computation.'
    )

    doc.add_heading('Kalman Rate Tracker', level=3)
    doc.add_paragraph(
        'A scalar Kalman filter fuses spectral and adaptive peak estimates '
        'within each 30-second window. The filter state is the current rate '
        'estimate in Hz. Process noise Q encodes the maximum physiological '
        'rate-of-change: 2 br/min per epoch for respiratory, 5 BPM per epoch '
        'for cardiac. Measurement noise R is set per-method based on benchmark '
        'MAEs. The filter was tuned for high reactivity (R scaled by 0.3, '
        'Q scaled by 2.0) so that it tracks genuine rate changes quickly while '
        'rejecting isolated outlier estimates.'
    )

    doc.add_heading('Multi-Channel Quality-Weighted Fusion', level=3)
    doc.add_paragraph(
        'Five channel configurations (CLE, CRE, CH, Avg, Diff) are processed '
        'independently. Per-window quality scores -- based on in-band SNR, '
        'spectral concentration, ACF prominence, motion power, and method '
        'agreement -- weight each channel\'s contribution to the fused estimate. '
        'For respiratory, each channel runs the full Kalman pipeline (spectral + '
        'adaptive peaks); for cardiac, each channel provides a Hilbert '
        'instantaneous frequency estimate directly.'
    )

    doc.add_heading('Temporal Smoothing and k-Scaling', level=3)
    doc.add_paragraph(
        'A median filter (width 7 epochs = 3.5 minutes) is applied to the '
        'fused rate trace to suppress residual jitter. The smoothing window '
        'width is a tunable parameter. Finally, a per-session k-factor '
        '(median ratio of CAP estimate to PSG ground truth) corrects systematic '
        'over- or under-counting. For cross-subject generalization, leave-one-'
        'subject-out (LOSO) k-calibration uses the median k from the remaining '
        '5 subjects.'
    )

    doc.add_heading('Final Pipeline Summary', level=3)
    doc.add_paragraph(
        'Respiratory: spectral + adaptive_peaks (per channel) --> reactive '
        'Kalman filter (per channel) --> quality-weighted multi-channel fusion '
        '--> median temporal smoothing --> k-scaling.'
    )
    doc.add_paragraph(
        'Cardiac: Hilbert instantaneous frequency (per channel) --> '
        'quality-weighted multi-channel fusion --> median temporal smoothing '
        '--> k-scaling.'
    )


def write_hybrid_results(doc: Document):
    doc.add_heading('Results', level=2)

    # --- Adaptive peaks benchmark ---
    doc.add_heading('Adaptive Peak Detector Benchmark', level=3)
    doc.add_paragraph(
        'The adaptive peak detector was benchmarked against five existing methods '
        'across all 12 sessions without k-scaling (Table 1). For respiratory rate, '
        'adaptive_peaks achieved 2.36 br/min MAE (vs. peaks 3.20, hilbert 2.50, '
        'spectral 1.92). For cardiac, adaptive_peaks achieved 26.51 BPM MAE '
        '(vs. peaks 48.53, hilbert 38.06, spectral 27.50). The adaptive detector '
        'is the best peak-based method for both bands: spectral-guided minimum '
        'distance eliminates most double-counting, and IPI validation rejects '
        'noisy windows.'
    )

    # --- Kalman tracker ---
    doc.add_heading('Kalman Tracker', level=3)
    doc.add_paragraph(
        'Fusing spectral and adaptive_peaks through the Kalman filter reduced '
        'cardiac MAE from 27.50 BPM (spectral alone) and 26.51 BPM (adaptive '
        'alone) to 21.22 BPM -- a 20% reduction over the best raw input. The '
        'RMSE improvement was even larger (32%), indicating that the Kalman '
        'filter primarily suppresses epoch-to-epoch jitter. For respiratory, '
        'the Kalman output matched spectral MAE (1.90 vs. 1.92 br/min) while '
        'adding temporal smoothness.'
    )

    add_figure(doc, PHASE1_DIR / 'aggregate_mae_comparison.png',
        'Figure 6. Kalman tracker aggregate comparison (no k). Spectral, '
        'adaptive peaks, and Kalman fused across 12 sessions. Cardiac: '
        'Kalman 21.2 BPM vs. spectral 27.5 (20% reduction).',
        width=6.0)

    # --- k-calibrated evaluation ---
    doc.add_heading('Per-Session k-Calibration', level=3)
    doc.add_paragraph(
        'With per-session k-calibration, the Kalman pipeline achieved 1.61 br/min '
        'respiratory MAE -- a 37% improvement over the baseline peaks/k (2.58 '
        'br/min), winning on all 12 sessions (Wilcoxon p=0.0002). However, for '
        'cardiac rate, the baseline hilbert/k (4.84 BPM) outperformed Kalman/k '
        '(8.67 BPM). The Hilbert instantaneous frequency captures BCG cardiac '
        'waveform physics more naturally: its per-session k (~1.67) corrects a '
        'consistent harmonic relationship, while the Kalman\'s spectral+adaptive '
        'inputs have a lower, less stable k (~1.34).'
    )

    add_figure(doc, PHASE3_DIR / 'session_comparison_resp.png',
        'Figure 7. Respiratory MAE per session: Kalman/k (red) vs. baseline '
        'peaks/k (blue). Kalman/k wins on all 12 sessions.',
        width=6.0)

    add_figure(doc, PHASE3_DIR / 'session_comparison_card.png',
        'Figure 8. Cardiac MAE per session: Kalman/k (red) vs. baseline '
        'hilbert/k (blue). Baseline hilbert/k wins on all 12 sessions.',
        width=6.0)

    # --- Multi-channel ---
    doc.add_heading('Multi-Channel Fusion', level=3)
    doc.add_paragraph(
        'Running the Kalman pipeline independently on 5 channels and fusing '
        'with quality weights yielded modest improvement over the single best '
        'channel (CLE-CRE diff): respiratory 1.82 vs. 1.90 br/min (4%), '
        'cardiac 17.74 vs. 21.22 BPM (16%). An oracle analysis selecting the '
        'best channel per window with hindsight achieved 1.21 br/min resp and '
        '8.63 BPM cardiac, showing substantial headroom for improved channel '
        'selection (36% resp, 59% cardiac above quality-weighted fusion).'
    )

    add_figure(doc, PHASE2_DIR / 'aggregate_comparison.png',
        'Figure 9. Multi-channel fusion: single-channel (diff) vs. '
        'quality-weighted vs. oracle. No k-scaling.',
        width=6.0)

    # --- Best pipeline aggregate ---
    doc.add_heading('Best Combined Pipeline', level=3)
    doc.add_paragraph(
        'The final pipeline combines Kalman (resp) + Hilbert (cardiac), '
        'multi-channel fusion, temporal smoothing (median filter, width=7), '
        'and per-session k-calibration. Table 2 shows the aggregate results '
        'across all 12 sessions.'
    )

    # Build results table
    csv_path = BEST_DIR / 'best_pipeline_results.csv'
    if csv_path.exists():
        df = pd.read_csv(csv_path)

        agg_rows = []
        for pipe in ['no_k', 'per_session_k', 'loso_k']:
            sub = df[df['pipeline'] == pipe]
            pipe_label = {'no_k': 'No k', 'per_session_k': 'Per-session k',
                          'loso_k': 'LOSO k'}[pipe]
            agg_rows.append({
                'Pipeline': pipe_label,
                'Resp MAE': f"{sub['resp_mae'].mean():.2f} +/- {sub['resp_mae'].std():.2f}",
                'Card MAE': f"{sub['card_mae'].mean():.2f} +/- {sub['card_mae'].std():.2f}",
                'Resp r': f"{sub['resp_r'].mean():.3f}",
                'Card r': f"{sub['card_r'].mean():.3f}",
            })
        agg_df = pd.DataFrame(agg_rows)
        add_table(doc, agg_df,
            'Table 2. Aggregate accuracy (mean +/- std across 12 sessions). '
            'Resp MAE in br/min, Card MAE in BPM. r = mean per-session '
            'Pearson correlation with PSG ground truth.')

    add_figure(doc, BEST_DIR / 'aggregate_comparison.png',
        'Figure 10. Aggregate MAE comparison: no k, per-session k, and '
        'LOSO k. Left: respiratory. Right: cardiac.',
        width=6.2)

    # --- Per-session breakdown ---
    doc.add_heading('Per-Session Results', level=3)

    if csv_path.exists():
        df = pd.read_csv(csv_path)
        ps = df[df['pipeline'] == 'per_session_k'].copy()
        ps = ps.sort_values('session')
        table_df = ps[['session', 'resp_mae', 'resp_r', 'k_resp',
                        'card_mae', 'card_r', 'k_card']].copy()
        table_df.columns = ['Session', 'Resp MAE', 'Resp r', 'k_resp',
                            'Card MAE', 'Card r', 'k_card']
        for col in ['Resp MAE', 'Card MAE']:
            table_df[col] = table_df[col].apply(lambda x: f'{x:.2f}')
        for col in ['Resp r', 'Card r']:
            table_df[col] = table_df[col].apply(lambda x: f'{x:.3f}')
        for col in ['k_resp', 'k_card']:
            table_df[col] = table_df[col].apply(lambda x: f'{x:.3f}')

        add_table(doc, table_df,
            'Table 3. Per-session results with per-session k-calibration. '
            'Resp MAE in br/min, Card MAE in BPM.')

    doc.add_paragraph(
        'The per-session k results show consistent respiratory accuracy across '
        'all 12 sessions (range 1.07-2.16 br/min). Cardiac accuracy is also '
        'consistent (range 2.54-6.56 BPM), with S2N2 as the worst case. The '
        'multi-channel Hilbert fusion eliminates the catastrophic S6 failures '
        'seen in the no-k consolidation pipeline, achieving 5.18 and 4.13 BPM '
        'for S6N1 and S6N2 respectively.'
    )

    # --- Per-stage ---
    doc.add_heading('Per-Stage Accuracy', level=3)

    add_figure(doc, BEST_DIR / 'per_stage_mae.png',
        'Figure 11. Per-stage MAE with per-session k. Left: respiratory. '
        'Right: cardiac. Sample sizes shown per stage.',
        width=6.2)

    # --- Bland-Altman aggregate ---
    doc.add_heading('Bland-Altman Agreement', level=3)

    add_figure(doc, PHASE3_DIR / 'bland_altman_aggregate_resp.png',
        'Figure 12. Bland-Altman: Kalman/k respiratory rate vs PSG, '
        'pooled across all 12 sessions.',
        width=5.5)

    add_figure(doc, PHASE3_DIR / 'bland_altman_aggregate_card.png',
        'Figure 13. Bland-Altman: Kalman/k cardiac rate vs PSG, '
        'pooled across all 12 sessions.',
        width=5.5)

    # --- Per-session time series ---
    doc.add_heading('Per-Session Time Series', level=3)
    doc.add_paragraph(
        'Figures 14-25 show the full-night rate traces for all 12 sessions. '
        'Each plot shows the ground truth (black), raw fused estimate (gray), '
        'and the final k-scaled smoothed estimate (red for respiratory, blue '
        'for cardiac), with sleep stage annotation along the bottom.'
    )

    for i, sess in enumerate(SESSIONS):
        fig_path = BEST_DIR / f'best_{sess}.png'
        fig_num = 14 + i
        add_figure(doc, fig_path,
            f'Figure {fig_num}. {sess} -- best pipeline time series '
            f'(respiratory + cardiac with per-session k).',
            width=6.5)

    # --- Per-session Bland-Altman ---
    doc.add_heading('Per-Session Bland-Altman', level=3)
    doc.add_paragraph(
        'Figures 26-37 show Bland-Altman agreement plots for each session.'
    )

    for i, sess in enumerate(SESSIONS):
        fig_path = BEST_DIR / f'bland_altman_{sess}.png'
        fig_num = 26 + i
        add_figure(doc, fig_path,
            f'Figure {fig_num}. {sess} -- Bland-Altman agreement '
            f'(respiratory left, cardiac right).',
            width=6.2)


# ==============================================================================
# DISCUSSION
# ==============================================================================

def write_discussion(doc: Document):
    doc.add_heading('Discussion', level=1)

    doc.add_heading('Pipeline Comparison', level=2)

    comp_data = {
        'Pipeline': [
            'Peaks/k baseline',
            'Consolidation: fused+Viterbi (no k)',
            'Hybrid: Kalman resp + hilbert/k cardiac',
            'Hybrid: LOSO k (cross-subject)',
        ],
        'Resp MAE': ['2.58', '1.50', '1.49', '1.95'],
        'Card MAE': ['4.84', '6.60', '4.11', '5.41'],
        'Requires PSG': ['Yes', 'No', 'Yes', 'Yes (other subjects)'],
    }
    comp_df = pd.DataFrame(comp_data)
    add_table(doc, comp_df,
        'Table 4. Summary comparison of all rate estimation pipelines. '
        'Resp MAE in br/min, Card MAE in BPM. Hybrid pipeline achieves '
        'best results in both bands.')

    doc.add_heading('Respiratory Rate', level=2)
    doc.add_paragraph(
        'The hybrid pipeline achieves 1.49 br/min MAE with per-session k, '
        'a 42% improvement over the original peaks/k baseline (2.58 br/min). '
        'Three factors contribute: (1) the adaptive peak detector eliminates '
        'most double-counting through spectral-guided minimum distance; '
        '(2) the reactive Kalman filter fuses spectral and adaptive peak '
        'estimates, suppressing jitter while tracking genuine rate changes; '
        '(3) multi-channel fusion provides a further 4% improvement by '
        'quality-weighting across five sensor configurations. The respiratory '
        'k values are near 1.0 (range 0.90-1.18), meaning the Kalman pipeline '
        'nearly eliminates the systematic overcount that k was designed to correct.'
    )

    doc.add_heading('Cardiac Rate', level=2)
    doc.add_paragraph(
        'For cardiac rate, the hybrid pipeline achieves 4.11 BPM MAE -- a 15% '
        'improvement over the previous best (hilbert/k on single channel, '
        '4.84 BPM). The improvement comes entirely from multi-channel fusion: '
        'the Hilbert method applied independently to five channels, '
        'quality-weighted, and temporally smoothed, outperforms any single '
        'channel. The cardiac k values (range 1.45-1.85, median 1.65) reflect '
        'the consistent BCG harmonic overcounting that Hilbert detects.'
    )
    doc.add_paragraph(
        'The Kalman pipeline was not used for cardiac because the spectral '
        'and adaptive peak inputs have a fundamentally different harmonic '
        'relationship with the BCG waveform than Hilbert instantaneous '
        'frequency. The Kalman cardiac k (~1.34) is lower and less stable '
        'than the Hilbert k (~1.67), resulting in worse k-corrected accuracy '
        '(8.67 vs 4.84 BPM). This suggests that Hilbert captures a physical '
        'property of the BCG -- likely the ratio of systolic to diastolic '
        'components -- that spectral peak frequency does not.'
    )

    doc.add_heading('Cross-Subject Generalization', level=2)
    doc.add_paragraph(
        'With LOSO k-calibration, accuracy degrades modestly: respiratory '
        '1.95 br/min (vs. 1.49 per-session), cardiac 5.41 BPM (vs. 4.11 '
        'per-session). The respiratory degradation is larger because k varies '
        'more across subjects (0.90-1.18 vs. cardiac 1.45-1.85). For '
        'deployment without PSG, the LOSO results represent realistic '
        'cross-subject performance using a population-level k.'
    )

    doc.add_heading('Limitations', level=2)
    doc.add_paragraph(
        'First, per-window correlation with ground truth remains low '
        '(mean r ~ 0.18 respiratory, -0.19 cardiac), indicating that the '
        'pipeline tracks average rates well but does not follow window-to-window '
        'fluctuations reliably. Second, the k-factor still requires PSG for '
        'calibration; while LOSO provides cross-subject generalization, a truly '
        'PSG-free cardiac pipeline with acceptable accuracy remains elusive. '
        'Third, the smoothing parameters (median filter width, Kalman R and Q '
        'scales) were selected by informal tuning rather than systematic '
        'optimization, and could likely be improved.'
    )


# ==============================================================================
# Main
# ==============================================================================

def main():
    doc = Document()
    setup_styles(doc)

    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        'Rate Estimation from Capacitive Temple Sensors:\n'
        'Multi-Channel Consolidation and Hybrid Pipeline'
    )
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.name = 'Times New Roman'

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('6 subjects, 12 overnight recordings, PSG ground truth')
    run.font.size = Pt(10)
    run.font.italic = True
    doc.add_paragraph()

    # Part 1: consolidation
    write_consolidation_methods(doc)
    write_consolidation_results(doc)

    # Part 2: hybrid pipeline
    write_hybrid_methods(doc)
    write_hybrid_results(doc)

    # Discussion
    write_discussion(doc)

    doc.save(str(OUT))
    print(f'Saved to {OUT}')

    # Count figures
    n_figs = 0
    n_figs += 5   # consolidation figures
    n_figs += 4   # hybrid aggregate figures
    n_figs += 2   # Bland-Altman aggregate
    n_figs += 12  # per-session time series
    n_figs += 12  # per-session Bland-Altman
    print(f'  {n_figs} figures, 4 tables')
    print(f'  Sections: Part 1 (consolidation), Part 2 (hybrid pipeline), Discussion')


if __name__ == '__main__':
    main()
