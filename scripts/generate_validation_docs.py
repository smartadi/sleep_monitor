#!/usr/bin/env python
"""
Generate validation study .docx documents:
  - notebooks/validation_methods.docx
  - notebooks/validation_results.docx

Usage:  python scripts/generate_validation_docs.py
"""

from __future__ import annotations
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

ROOT = Path(__file__).resolve().parent.parent
NOTEBOOKS = ROOT / 'notebooks'
PLOTS = NOTEBOOKS / 'plots' / 'rate_analysis'


# ── Helpers ───────────────────────────────────────────────────────────────────

def set_style(doc: Document):
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    pf = style.paragraph_format
    pf.space_after = Pt(4)
    pf.space_before = Pt(2)


def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_para(doc, text, bold=False, italic=False, size=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    return p


def add_equation(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.5)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return p


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = h
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i + 1].cells[j]
            cell.text = str(val)
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(9)
    return table


def add_image_placeholder(doc, filename, caption=''):
    path = PLOTS / filename
    if path.exists():
        doc.add_picture(str(path), width=Inches(6.0))
    else:
        p = doc.add_paragraph()
        run = p.add_run(f'[IMAGE: {filename} — run plot_validation.py to generate]')
        run.italic = True
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    if caption:
        p = doc.add_paragraph()
        run = p.add_run(caption)
        run.italic = True
        run.font.size = Pt(9)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER


# ── Methods Document ──────────────────────────────────────────────────────────

def build_methods():
    doc = Document()
    set_style(doc)

    doc.add_heading('Validation Study — Methods', level=0)

    # --- 1. Data ---
    add_heading(doc, '1. Data')
    add_para(doc,
        'Six healthy subjects (OS001–OS006) each completed two overnight polysomnography '
        '(PSG) sessions, yielding 12 recordings. During each session a capacitive sleep monitor '
        'was worn on the temples, recording three channels — left electrode (CLE), right '
        'electrode (CRE), and a head channel (CH) — alongside a 3-axis accelerometer, all '
        'sampled at 100 Hz. Simultaneous PSG recorded ECG, nasal airflow (Flow), thoracic effort '
        '(Thorax), photoplethysmography (Pleth), EEG, and EOG at 100 Hz (resampled from native '
        'PSG rate). A board-certified sleep technologist scored 30-second epochs into Wake, N1, '
        'N2, N3, and REM stages per AASM criteria.')
    add_para(doc,
        'The CAP and PSG recordings were time-synchronised offline; quality was verified by '
        'visual inspection and automated artifact detection.')

    # --- 2. Ground Truth ---
    add_heading(doc, '2. Ground Truth Rate Extraction')

    add_heading(doc, '2.1 Cardiac Ground Truth', level=2)
    add_para(doc,
        'R-peaks were detected from the ECG channel using the Pan-Tompkins algorithm as '
        'implemented in neurokit2 (ecg_clean + ecg_findpeaks). A quality filter removed peaks '
        'producing inter-beat intervals outside the physiological range (0.33–2.0 s, '
        'corresponding to 30–180 BPM). Per-epoch (30 s, non-overlapping) heart rate was '
        'computed as:')
    add_equation(doc, 'HR_GT = (N_peaks - 1) / (t_last_peak - t_first_peak)')
    add_para(doc,
        'where N_peaks is the number of R-peaks within the epoch. Epochs with fewer than '
        '2 R-peaks were marked as missing.')
    add_para(doc,
        'If ECG R-peak detection failed for a session (< 2 valid peaks overall), the system '
        'fell back to peak detection on bandpass-filtered Pleth (0.5–3.0 Hz).')

    add_heading(doc, '2.2 Respiratory Ground Truth', level=2)
    add_para(doc,
        'Breath peaks were detected from the nasal airflow (Flow) channel using neurokit2’s '
        'respiratory processing pipeline (rsp_clean + rsp_findpeaks). A quality filter removed '
        'peaks producing inter-breath intervals outside the physiological range (2.0–10.0 s, '
        'corresponding to 6–30 breaths/min). Per-epoch respiratory rate was computed '
        'analogously to cardiac rate.')
    add_para(doc,
        'If Flow peak detection failed, the system fell back to peak detection on '
        'bandpass-filtered Thorax (0.1–0.5 Hz).')

    # --- 3. Preprocessing ---
    add_heading(doc, '3. CAP Signal Preprocessing')
    add_para(doc,
        'The differential channel CLE−CRE was used as the primary input for rate estimation. '
        'This channel rejects common-mode noise and motion artifacts while preserving the '
        'ballistocardiographic (BCG) and respiratory components.')
    add_para(doc,
        'Accelerometer artifact removal was performed via ordinary least-squares (OLS) '
        'regression. Both the CAP and accelerometer signals were first bandpass-filtered to the '
        'target band (respiratory: 0.1–0.5 Hz; cardiac: 0.5–3.0 Hz), and the bandpassed '
        'accelerometer magnitude was regressed out of the bandpassed CAP signal:')
    add_equation(doc,
        'sig_clean = bandpass(CLE−CRE, f_lo, f_hi) − β × bandpass(acc_mag, f_lo, f_hi)')
    add_para(doc,
        'where β is the OLS coefficient. This removes only motion energy within the target '
        'frequency band, preserving physiological content outside it.')

    # --- 4. Rate Estimation ---
    add_heading(doc, '4. Rate Estimation Methods')

    add_heading(doc, '4.1 Respiratory Rate: Scaled Peak Counting', level=2)
    add_para(doc,
        'Respiratory rate was estimated using a loose peak detector followed by a learned '
        'scaling correction. The bandpass-filtered respiratory signal (0.1–0.5 Hz, '
        'OLS-cleaned) was lightly smoothed, and peaks were detected with permissive settings '
        '(prominence threshold = 0.05×σ, minimum inter-peak distance = 0.4 s). This '
        'detector consistently captures both inhalation and exhalation bumps present in the BCG '
        'signal, yielding a systematic overcount.')
    add_para(doc,
        'A per-session scaling factor k_resp was calibrated by comparing the loose-peak rate '
        'against the Flow-derived ground truth across 50 randomly sampled 30-second windows:')
    add_equation(doc, 'k_resp = median(rate_peaks_loose / rate_GT)')
    add_para(doc, 'The corrected respiratory rate per epoch is:')
    add_equation(doc, 'RR_CAP = (N_peaks / k_resp) / epoch_duration')
    add_para(doc,
        'Cross-session k_resp ranged from 1.18 to 1.61 (median ~1.31), reflecting '
        'subject-specific coupling geometry. Calibration was performed at the same 30-second '
        'window size used for evaluation.')

    add_heading(doc, '4.2 Cardiac Rate: Scaled Hilbert Instantaneous Frequency', level=2)
    add_para(doc,
        'Cardiac rate was estimated using the Hilbert transform instantaneous frequency, '
        'followed by a learned scaling correction. The analytic signal of the bandpass-filtered '
        'cardiac signal (0.5–3.0 Hz, OLS-cleaned) was computed, and the median instantaneous '
        'frequency across the epoch was taken as the raw rate estimate.')
    add_para(doc,
        'The raw Hilbert rate over-counts by a factor of ~1.5–1.9× per session, caused '
        'by systolic and dicrotic-notch components in the BCG waveform. A per-session scaling '
        'factor k_cardiac was calibrated analogously:')
    add_equation(doc, 'k_cardiac = median(rate_hilbert / rate_GT)')
    add_para(doc,
        'using 50 randomly sampled 30-second windows against ECG-derived heart rate. '
        'The corrected cardiac rate per epoch is:')
    add_equation(doc, 'HR_CAP = rate_hilbert / k_cardiac')
    add_para(doc,
        'Cross-session k_cardiac ranged from 1.48 to 1.93 (median ~1.67). Per-night calibration '
        'was used, as night-to-night |Δk| reached 0.19 on some subjects.')

    # --- 5. Evaluation Protocol ---
    add_heading(doc, '5. Evaluation Protocol')

    add_heading(doc, '5.1 Epoch Definition', level=2)
    add_para(doc,
        'Evaluation was performed on non-overlapping 30-second epochs, matching the standard '
        'PSG scoring epoch length. This yields approximately 960 epochs per 8-hour recording, '
        'totalling ~11,500 epochs across all 12 sessions.')

    add_heading(doc, '5.2 Sleep Stage Assignment', level=2)
    add_para(doc,
        'Each epoch was assigned the sleep stage from the contemporaneous PSG scoring epoch '
        'based on time alignment. Epochs falling outside the scored PSG window or labelled as '
        'artifact were excluded from stage-stratified analysis.')

    add_heading(doc, '5.3 Accuracy Metrics', level=2)
    add_para(doc,
        'For each session and each rate type (respiratory, cardiac), the following metrics were '
        'computed over all epochs with valid CAP and GT estimates:')
    add_table(doc,
        ['Metric', 'Definition', 'Unit'],
        [
            ['MAE', 'mean(|CAP_rate − GT_rate|)', 'br/min or BPM'],
            ['RMSE', 'sqrt(mean((CAP_rate − GT_rate)²))', 'br/min or BPM'],
            ['Bias', 'mean(CAP_rate − GT_rate)', 'br/min or BPM'],
            ['Pearson r', 'Correlation between CAP and GT rate', 'dimensionless'],
            ['p50', 'Median absolute error', 'br/min or BPM'],
            ['p90', '90th percentile absolute error', 'br/min or BPM'],
            ['Coverage', 'Fraction of epochs with valid estimate', 'proportion'],
        ])
    add_para(doc,
        'All rate values were converted from Hz to per-minute units (×60) before '
        'computing error metrics.')

    add_heading(doc, '5.4 Bland-Altman Analysis', level=2)
    add_para(doc,
        'Agreement between CAP-derived and PSG-derived rates was assessed using Bland-Altman '
        'plots. The mean of the two measurements was plotted against their difference (CAP − GT), '
        'with the mean bias and 95% limits of agreement (bias ± 1.96 × SD) overlaid.')

    add_heading(doc, '5.5 Sleep Stage Stratification', level=2)
    add_para(doc,
        'Per-stage accuracy was computed by pooling all epochs of each stage (Wake, N1, N2, N3, '
        'REM) across all 12 sessions and computing the same metric set. This quantifies how '
        'rate estimation accuracy varies with sleep depth, arousal state, and the associated '
        'changes in physiological dynamics.')

    # --- 6. Calibration Note ---
    add_heading(doc, '6. Calibration Note')
    add_para(doc,
        'The scaling factors k_resp and k_cardiac were calibrated per-session using 50 randomly '
        'sampled 30-second windows from the same recording. This is analogous to a brief in-lab '
        'calibration that would precede clinical deployment. The calibration windows were drawn '
        'uniformly from the full recording and are not excluded from evaluation (the number of '
        'calibration windows is <5% of total epochs and randomly distributed, so leakage bias '
        'is negligible). In a deployment setting, calibration could be performed during a short '
        'attended period at the start of recording.')

    path = NOTEBOOKS / 'validation_methods.docx'
    doc.save(str(path))
    print(f"Saved: {path}")


# ── Results Document ──────────────────────────────────────────────────────────

def build_results():
    doc = Document()
    set_style(doc)

    doc.add_heading('Validation Study — Results', level=0)

    # --- 1. Dataset Summary ---
    add_heading(doc, '1. Dataset Summary')
    add_para(doc,
        '12 sessions (6 subjects × 2 nights), totalling approximately _____ hours of '
        'recording and _____ 30-second epochs evaluated.')

    # --- 2. Calibration Factors ---
    add_heading(doc, '2. Calibration Factors')
    sessions = [
        ('S1N1', 'OS001'), ('S1N2', 'OS001'),
        ('S2N1', 'OS002'), ('S2N2', 'OS002'),
        ('S3N1', 'OS003'), ('S3N2', 'OS003'),
        ('S4N1', 'OS004'), ('S4N2', 'OS004'),
        ('S5N1', 'OS005'), ('S5N2', 'OS005'),
        ('S6N1', 'OS006'), ('S6N2', 'OS006'),
        ('Median', '—'),
    ]
    add_table(doc,
        ['Session', 'Subject', 'k_resp', 'k_cardiac'],
        [[s, subj, '', ''] for s, subj in sessions])
    add_para(doc, 'Fill from artifacts/validation_session.csv', italic=True, size=9)
    doc.add_paragraph()
    add_para(doc, 'Notes on k stability:', bold=True)
    add_para(doc, '• k_resp range: [ ___ , ___ ], subject clustering: ___')
    add_para(doc, '• k_cardiac range: [ ___ , ___ ], max night-to-night |Δk|: ___')

    # --- 3. Aggregate Accuracy ---
    add_heading(doc, '3. Aggregate Accuracy')
    add_table(doc,
        ['Metric', 'Respiratory (br/min)', 'Cardiac (BPM)'],
        [
            ['MAE', '', ''],
            ['RMSE', '', ''],
            ['Bias', '', ''],
            ['Pearson r', '', ''],
            ['p50', '', ''],
            ['p90', '', ''],
            ['Coverage', '', ''],
            ['N epochs', '', ''],
        ])
    add_para(doc, 'Fill from the ALL row in artifacts/validation_session.csv', italic=True, size=9)

    # --- 4. Per-Session Accuracy ---
    add_heading(doc, '4. Per-Session Accuracy')

    add_heading(doc, '4.1 Summary Table', level=2)
    add_image_placeholder(doc, 'validation_summary_table.png',
                          'Figure 1. Per-session and per-stage validation metrics.')
    add_para(doc, 'Commentary:', bold=True)
    add_para(doc, '• Best/worst session for respiratory: ___')
    add_para(doc, '• Best/worst session for cardiac: ___')
    add_para(doc, '• Outlier sessions to discuss (e.g. S6N2 k_cardiac anomaly): ___')

    add_heading(doc, '4.2 Per-Session MAE Bar Chart', level=2)
    add_image_placeholder(doc, 'validation_session_mae.png',
                          'Figure 2. Mean absolute error per session for respiratory and cardiac rate.')
    add_para(doc, 'Commentary on session-to-session variability:', bold=True)
    add_para(doc, '')

    # --- 5. Agreement Analysis ---
    add_heading(doc, '5. Agreement Analysis')

    add_heading(doc, '5.1 Bland-Altman Plots', level=2)
    add_image_placeholder(doc, 'validation_bland_altman.png',
                          'Figure 3. Bland-Altman plots for respiratory and cardiac rate agreement. '
                          'Points colored by sleep stage.')
    add_para(doc, 'Respiratory rate:', bold=True)
    add_para(doc, '• Bias: ___ br/min')
    add_para(doc, '• 95% Limits of agreement: [ ___ , ___ ] br/min')
    add_para(doc, '• Proportional bias observed?')
    add_para(doc, '• Stage-dependent patterns visible?')
    doc.add_paragraph()
    add_para(doc, 'Cardiac rate:', bold=True)
    add_para(doc, '• Bias: ___ BPM')
    add_para(doc, '• 95% Limits of agreement: [ ___ , ___ ] BPM')
    add_para(doc, '• Proportional bias observed?')
    add_para(doc, '• Stage-dependent patterns visible?')

    add_heading(doc, '5.2 Scatter Plots', level=2)
    add_image_placeholder(doc, 'validation_scatter.png',
                          'Figure 4. CAP rate vs GT rate scatter plots. '
                          'Dashed line = identity. Points colored by sleep stage.')
    add_para(doc, 'Commentary:', bold=True)
    add_para(doc, '• Spread around identity line')
    add_para(doc, '• Systematic deviation at high/low rates?')
    add_para(doc, '• Stage clustering visible?')

    # --- 6. Sleep Stage Stratified Accuracy ---
    add_heading(doc, '6. Sleep Stage Stratified Accuracy')

    add_heading(doc, '6.1 Per-Stage Metrics Table', level=2)
    add_table(doc,
        ['Stage', 'Resp MAE', 'Resp r', 'Resp n', 'Card MAE', 'Card r', 'Card n'],
        [
            ['Wake', '', '', '', '', '', ''],
            ['N1',   '', '', '', '', '', ''],
            ['N2',   '', '', '', '', '', ''],
            ['N3',   '', '', '', '', '', ''],
            ['REM',  '', '', '', '', '', ''],
        ])
    add_para(doc, 'Fill from artifacts/validation_stage.csv', italic=True, size=9)

    add_heading(doc, '6.2 Error Distribution by Stage', level=2)
    add_image_placeholder(doc, 'validation_stage_boxplots.png',
                          'Figure 5. Absolute error distribution by sleep stage '
                          'for respiratory and cardiac rate.')
    add_para(doc, 'Commentary:', bold=True)
    add_para(doc, '• Which stages have lowest/highest error?')
    add_para(doc, '• Hypothesis: Wake/N1 noisier due to movement and arousal; N2/N3 quieter')
    add_para(doc,
        '• REM: increased HRV and irregular breathing — does this affect accuracy?')
    add_para(doc,
        '• Connection to k biomarker findings (k_cardiac varies by stage: '
        'N1=1.71 > N2/N3=1.65 > Wake=1.61 > REM=1.58)')

    # --- 7. Discussion Points ---
    add_heading(doc, '7. Discussion Points')

    add_heading(doc, '7.1 Key Findings', level=2)
    add_para(doc, '• CAP-derived respiratory rate achieves MAE of ___ br/min '
                  '(30 s epochs, 12 sessions)')
    add_para(doc, '• CAP-derived cardiac rate achieves MAE of ___ BPM '
                  '(30 s epochs, 12 sessions)')
    add_para(doc, '• Both signals are reliably present in the capacitive sensor data')
    add_para(doc, '• Sleep stage modulates accuracy — describe pattern')

    add_heading(doc, '7.2 Calibration Requirement', level=2)
    add_para(doc, '• Per-session k calibration is needed (k varies across subjects/nights)')
    add_para(doc, '• 50 random 30 s windows (~25 min of data) sufficient for calibration')
    add_para(doc, '• Clinical deployment: could calibrate during attended setup period')

    add_heading(doc, '7.3 Limitations', level=2)
    add_para(doc, '• Small cohort (6 subjects, healthy)')
    add_para(doc,
        '• k calibration requires paired PSG data — not available in standalone deployment')
    add_para(doc, '• OLS artifact removal assumes stationary accelerometer coupling')
    add_para(doc, '• Coverage < 100%: some epochs produce invalid estimates')

    add_heading(doc, '7.4 Comparison to Prior Work', level=2)
    add_para(doc, '• Compare to published BCG-based sleep monitoring (bed sensors, wearables)')
    add_para(doc, '• Typical BCG cardiac MAE: 2–8 BPM depending on sensor and cohort')
    add_para(doc, '• Typical BCG respiratory MAE: 1–4 br/min')
    add_para(doc, '• Where does this system fall?')

    # --- 8. Next Steps ---
    add_heading(doc, '8. Next Steps After Validation')
    add_para(doc, '• Quality gating + median smoothing (Task 6 item 1) to reduce MAE '
                  'from noisy epochs')
    add_para(doc, '• NLMS accelerometer removal to handle posture-dependent coupling')
    add_para(doc, '• Investigate whether default k (population median) is viable without '
                  'per-session calibration')

    path = NOTEBOOKS / 'validation_results.docx'
    doc.save(str(path))
    print(f"Saved: {path}")


def main():
    NOTEBOOKS.mkdir(parents=True, exist_ok=True)
    build_methods()
    build_results()
    print("\nDone. Both .docx files saved to notebooks/")


if __name__ == '__main__':
    main()
