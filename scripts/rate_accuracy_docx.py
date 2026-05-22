#!/usr/bin/env python
"""Generate DOCX report for rate accuracy analysis."""

from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt

PLOT_DIR = Path(__file__).resolve().parent.parent / 'notebooks' / 'plots' / 'rate_accuracy'
OUT = Path(__file__).resolve().parent.parent / 'artifacts' / 'Rate_Accuracy_Analysis.docx'

doc = Document()
style = doc.styles['Normal']
style.font.size = Pt(10)
style.font.name = 'Calibri'

doc.add_heading('Rate Accuracy Analysis', level=0)
doc.add_paragraph(
    'Full-night rate detection accuracy across 4 cap sensor channel combinations, '
    '12 sessions (6 subjects x 2 nights), stratified by sleep stage, apnea status, '
    'motion level, and electrode drift.'
)

# 1. Methods
doc.add_heading('1. Methods', level=1)
doc.add_paragraph(
    'Rate estimation was performed on non-overlapping 30-second epochs across all 12 overnight recordings. '
    'Four channel combinations were evaluated: avg = (CLE + CRE) / 2, diff = CLE - CRE, CLE alone, and CRE alone. '
    'All channels were bandpassed and accelerometer-artifact-removed (OLS) prior to rate estimation.'
)
doc.add_paragraph(
    'Respiratory rate: loose peak detection with per-session scaling factor k (rate_peaks_scaled_resp). '
    'Cardiac rate: Hilbert instantaneous frequency with per-session k (rate_hilbert_scaled_cardiac). '
    'k was calibrated per channel per session from 50 random 60-second windows against GT.'
)
doc.add_paragraph(
    'Ground truth: Flow (nasal airflow, neurokit2) for respiratory rate. '
    'ECG R-peaks (Pan-Tompkins) for cardiac rate in 10 sessions; Pleth peak detection '
    '(stricter settings, min_dist=0.4s) for S5N1 and S6N2 where ECG was unavailable.'
)
doc.add_paragraph(
    'Each epoch was tagged with: sleep stage (from PSG sleep profile), apnea status '
    '(Normal/Apnea/Hypopnea from PSG Flow annotations), accelerometer RMS (motion), '
    'and CLE/CRE raw signal mean values and epoch-to-epoch deltas (electrode drift).'
)

# 2. Overall Results
doc.add_heading('2. Overall Results', level=1)

table = doc.add_table(rows=6, cols=5, style='Light Grid Accent 1')
headers = ['Channel', 'Resp MAE (br/min)', 'Resp Bias', 'Card MAE (BPM)', 'Card Bias']
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h

data = [
    ('avg',  '2.78', '+0.82', '5.75', '-2.31'),
    ('diff', '2.68', '',      '5.55', ''),
    ('CLE',  '2.69', '',      '5.77', ''),
    ('CRE',  '2.62', '',      '5.42', ''),
    ('Oracle best', '0.23', '', '0.65', ''),
]
for ri, (ch, rm, rb, cm, cb) in enumerate(data):
    row = table.rows[ri + 1]
    row.cells[0].text = ch
    row.cells[1].text = rm
    row.cells[2].text = rb
    row.cells[3].text = cm
    row.cells[4].text = cb

doc.add_paragraph('')
doc.add_paragraph(
    'No single channel dominates. CRE has the lowest cardiac MAE (5.42 BPM) and respiratory MAE (2.62 br/min), '
    'but all four channels are within 0.4 BPM and 0.2 br/min of each other. '
    'The oracle best-channel selector (picking the lowest-error channel per epoch with GT knowledge) '
    'achieves 0.23 br/min resp and 0.65 BPM cardiac -- a 10-12x improvement over any fixed channel. '
    'This demonstrates that the best channel varies dramatically per epoch, and adaptive selection has '
    'enormous potential.'
)
doc.add_paragraph(
    'Oracle channel distribution: CRE wins 27% of epochs (resp and cardiac), followed by diff (23%), '
    'CLE (22%), and avg (21%). The near-uniform distribution confirms no single channel is reliably best.'
)

# 3. Per-Session Summary
doc.add_heading('3. Per-Session Variability', level=1)
doc.add_picture(str(PLOT_DIR / 'fig6_per_session_summary.png'), width=Inches(6))
doc.add_paragraph(
    'Figure 6: Per-session MAE for avg channel. Respiratory MAE ranges from 2.21 (S5N2) to 3.69 (S2N2) br/min. '
    'Cardiac MAE ranges from 3.81 (S5N2) to 8.42 (S6N2) BPM. S6N2 uses Pleth GT (ECG unavailable), '
    'which may contribute to higher cardiac error. S2N2 is the worst respiratory session, '
    'consistent with its high apnea burden (215 events).'
)

# 4. Overnight Tracking
doc.add_heading('4. Overnight Rate Tracking', level=1)
doc.add_picture(str(PLOT_DIR / 'fig1_overnight_rates.png'), width=Inches(6.5))
doc.add_paragraph(
    'Figure 1: Full-night rate time series for S1N1 (top, moderate apnea) and S2N2 (bottom, heavy apnea). '
    'CAP respiratory rate (green) tracks GT Flow (black) well during stable sleep. '
    'Cardiac rate (red vs black) shows systematic undercounting (negative bias of -2.3 BPM on average) '
    'but follows the GT trend. Major deviations coincide with apnea event markers in the hypnogram strip.'
)

# 5. Sleep Stage
doc.add_heading('5. Effect of Sleep Stage', level=1)
doc.add_picture(str(PLOT_DIR / 'fig2_error_by_stage.png'), width=Inches(6))

table2 = doc.add_table(rows=6, cols=4, style='Light Grid Accent 1')
for i, h in enumerate(['Stage', 'n epochs', 'Resp MAE', 'Card MAE']):
    table2.rows[0].cells[i].text = h
stage_data = [
    ('Wake',  '1,224', '2.77', '6.54'),
    ('N1',    '1,459', '2.64', '5.13'),
    ('N2',    '5,526', '2.80', '5.57'),
    ('N3',    '781',   '2.91', '6.45'),
    ('REM',   '223',   '2.79', '5.41'),
]
for ri, (st, n, rm, cm) in enumerate(stage_data):
    row = table2.rows[ri + 1]
    row.cells[0].text = st
    row.cells[1].text = n
    row.cells[2].text = rm
    row.cells[3].text = cm

doc.add_paragraph('')
doc.add_paragraph(
    'Figure 2: Respiratory error is remarkably stable across stages (2.64-2.91 br/min). '
    'Cardiac error is lowest during N1 (5.13) and REM (5.41), highest during Wake (6.54) and N3 (6.45). '
    'Wake cardiac error likely reflects increased motion artifacts. '
    'N3 cardiac error may reflect reduced pulse amplitude during deep sleep. '
    'The stage dependence of cardiac error is moderate -- no stage is catastrophically worse.'
)

# 6. Apnea
doc.add_heading('6. Effect of Apnea Events', level=1)
doc.add_picture(str(PLOT_DIR / 'fig3_error_by_apnea.png'), width=Inches(6))
doc.add_paragraph(
    'Figure 3: Apnea epochs show modestly higher respiratory error (3.10 vs 2.76 br/min for normal epochs), '
    'consistent with disrupted breathing patterns. Hypopnea epochs are intermediate (2.90 br/min). '
    'Cardiac error is slightly elevated during apnea (6.44 vs 5.73 BPM) but not dramatically. '
    'Surprisingly, hypopnea cardiac error (5.65 BPM) is comparable to normal epochs. '
    'The relatively modest impact of apnea on rate accuracy suggests the cap sensor continues to detect '
    'physiological signals even during respiratory events, though with reduced accuracy.'
)

# 7. Motion and Drift
doc.add_heading('7. Motion vs Electrode Drift', level=1)
doc.add_picture(str(PLOT_DIR / 'fig4_error_by_motion_and_drift.png'), width=Inches(6.5))

doc.add_paragraph(
    'Figure 4: This is a key finding. Accelerometer-based motion (top row) shows NO clear monotonic '
    'relationship with rate error -- in fact, the lowest-motion quartile (Q1) has HIGHER cardiac error '
    'than Q3/Q4. This is counterintuitive and suggests that OLS accelerometer removal is handling motion '
    'artifacts adequately.'
)
doc.add_paragraph(
    'In contrast, electrode drift (bottom row) -- the epoch-to-epoch change in raw CLE/CRE mean values -- '
    'is a strong predictor of error. Cardiac MAE rises from 3.84 BPM in Q1 (low drift) to 10.05 BPM in Q4 '
    '(high drift), a 2.6x increase. Respiratory MAE goes from 2.56 to 3.26 br/min. '
    'Electrode drift captures mask repositioning, sweat accumulation, and tissue compliance changes that '
    'are NOT detected by the accelerometer. This makes electrode drift a prime candidate for quality gating.'
)

# 8. Bland-Altman
doc.add_heading('8. Bland-Altman Analysis', level=1)
doc.add_picture(str(PLOT_DIR / 'fig5_bland_altman.png'), width=Inches(6.5))
doc.add_paragraph(
    'Figure 5: Bland-Altman plots for avg channel. Respiratory rate shows a positive bias of +0.82 br/min '
    '(slight overcounting). Cardiac rate shows a negative bias of -2.31 BPM (undercounting). '
    'Points are colored by sleep stage; no stage shows systematic outlier behavior. '
    'The spread is wider for cardiac, consistent with the higher MAE.'
)

# 9. Channel Comparison
doc.add_heading('9. Channel Comparison', level=1)
doc.add_picture(str(PLOT_DIR / 'fig7_channel_comparison.png'), width=Inches(6))
doc.add_paragraph(
    'Figure 7: All four channels perform similarly (MAE within 0.4 BPM cardiac, 0.2 br/min resp). '
    'The oracle best-channel line (dashed) shows the theoretical lower bound achievable with perfect '
    'per-epoch channel selection: 0.23 br/min resp, 0.65 BPM cardiac. '
    'This 10x gap between fixed and oracle channel selection is the strongest motivation for the '
    'GT-free adaptive channel selector planned in the next phase.'
)

# 10. Best Channel Distribution
doc.add_heading('10. Best Channel by Sleep Stage', level=1)
doc.add_picture(str(PLOT_DIR / 'fig8_best_channel_distribution.png'), width=Inches(6))
doc.add_paragraph(
    'Figure 8: The best channel is approximately uniformly distributed across all four options, '
    'with only mild stage dependence. CRE is slightly favored in N2 and N3 for cardiac, '
    'while diff is slightly favored during Wake for respiratory. '
    'The near-uniform distribution means simple heuristics (e.g., always use avg) will not capture '
    'the available improvement -- a per-epoch selector using signal quality features is needed.'
)

# 11. Key Findings
doc.add_heading('11. Key Findings and Next Steps', level=1)

findings = [
    ('1. Rate detection works across full nights with MAE of 2.78 br/min (resp) and 5.75 BPM (cardiac) '
     'using the avg channel with per-session k calibration.'),
    ('2. Electrode drift (CLE/CRE DC mean shifts) is a stronger predictor of rate error than motion. '
     'Cardiac MAE is 2.6x higher in the highest drift quartile (10.05 vs 3.84 BPM). Quality gating '
     'based on drift magnitude should be the first improvement to implement.'),
    ('3. No single channel is consistently best. Oracle per-epoch channel selection achieves '
     '0.23 br/min / 0.65 BPM -- a 10-12x improvement. This justifies building a GT-free channel selector.'),
    ('4. Sleep stage has moderate effect on cardiac error (Wake/N3 worse, N1/REM better) '
     'but minimal effect on respiratory error.'),
    ('5. Apnea events modestly increase respiratory error (+12%) but do not catastrophically degrade '
     'rate detection.'),
]
for f in findings:
    doc.add_paragraph(f, style='List Number')

doc.add_paragraph('')
doc.add_paragraph('Next steps:')
next_steps = [
    ('Quality gating: gate out epochs with high electrode drift (|delta_CLE| + |delta_CRE| in Q4), '
     'apply causal median smoothing to remaining epochs.'),
    ('GT-free channel selector: extract per-epoch cap-signal-only features (SNR, spectral clarity, '
     'L/R amplitude ratio, inter-channel coherence) and train classifier on oracle labels.'),
    ('Stage-aware k: use sleep stage to select different k values '
     '(validated in k biomarker analysis: k varies by stage).'),
]
for s in next_steps:
    doc.add_paragraph(s, style='List Bullet')

doc.save(str(OUT))
print(f'Saved {OUT}')
